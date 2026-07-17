from __future__ import annotations

import argparse
import asyncio
import calendar
import copy
import csv
import json
import math
import re
import sys
import unicodedata
from dataclasses import dataclass, field
from datetime import date, datetime
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Iterable, Sequence
from urllib.parse import quote_plus, urljoin, urlparse

import xlrd
import xlwt
import yaml
from docx import Document
from playwright.async_api import (
    Browser,
    BrowserContext,
    Download,
    Locator,
    Page,
    TimeoutError as PlaywrightTimeoutError,
    async_playwright,
)


DEFAULT_BASE_URL = "https://myrosmol.ru"
MODE_DOWNLOAD = "download"
MODE_UPLOAD = "upload"
MODE_ALL = "all"
MODE_PARSE = "parse"

PROJECT_TEMPLATE_DEFAULT = (
    "2026: Шаблон проекта конкурса Росмолодёжь.Гранты среди физических лиц"
)

LOGIN_BUTTON_PATTERN = re.compile(
    r"^\s*(Войти|Вход|Авторизоваться|Личный кабинет)\s*$", re.I
)
PASSWORD_MODE_PATTERN = re.compile(
    r"Войти по паролю|Вход по паролю|логин и пароль", re.I
)
SAVE_BUTTON_PATTERN = re.compile(
    r"^\s*(Сохранить черновик|Сохранить проект|Сохранить)\s*$", re.I
)
ADD_PROJECT_PATTERN = re.compile(r"Добавить проект|Создать проект", re.I)
CREATE_DRAFT_PATTERN = re.compile(r"Создать черновик", re.I)
DOWNLOAD_PROJECT_PATTERN = re.compile(
    r"Скачать проект|Выгрузить проект|Скачать заявку|Выгрузить заявку", re.I
)

PERSONAL_FIELDS = {
    "имя",
    "e mail",
    "email",
    "телефон",
    "дата рождения",
    "опыт автора проекта",
    "описание функционала автора проекта",
    "адрес регистрации автора проекта",
    "добавить резюме поле для файла",
    "добавить резюме",
    "видео визитка ссылка на ролик на любом видеохостинге",
    "регион нахождения автора проекта выборное поле с точностью до адреса можно добавить несколько пунктов",
    "регион нахождения автора проекта",
}

SKIPPED_SECTIONS = {"команда", "медиа"}

EXPENSE_CATEGORIES = (
    "Расходы на создание и/или техническую поддержку сайта",
    "Расходы на телефонную связь, мобильную связь, информационно-телекоммуникационной сети «Интернет», почтовые расходы",
    "Расходы на канцелярские принадлежности",
    "Расходы на издательско-полиграфические услуги, в т.ч. изготовление макета, разработка дизайна",
    "Расходы на подарки, сувенирную продукцию",
    "Расходы на проживание и питание",
    "Транспортные расходы (приобретение авиа- и железнодорожных билетов, горюче-смазочных материалов, услуги по перевозке пассажиров)",
    "Расходы на аренду помещения",
    "Расходы на аренду оборудования",
    "Расходы на информационные услуги (размещение информации о проекте в средствах массовой информации)",
    "Расходы на закупку оборудования",
    "Закупка расходных материалов и оплата услуг, необходимых для реализации проекта",
    "Расходы на покупку и/или создание программного обеспечения",
)


class AutomationError(RuntimeError):
    """Ошибка одной строки или настройки с понятным сообщением для пользователя."""


async def ask_error_action(
    step_name: str,
    error: Exception,
    *,
    skip_label: str = "пропустить только этот шаг",
) -> str:
    """Запрашивает действие после ошибки: повторить, пропустить шаг или остановить."""
    print(f"\nОШИБКА НА ЭТАПЕ: {step_name}")
    print(f"Причина: {error}")
    print("1 - повторить текущий шаг")
    print(f"2 - {skip_label}")
    print("3 - остановить программу")
    while True:
        try:
            raw = (
                await asyncio.to_thread(
                    input,
                    "Выберите действие [Enter = повторить]: ",
                )
            ).strip().casefold()
        except EOFError as exc:
            raise AutomationError(
                f"Не удалось запросить действие после ошибки на этапе «{step_name}»"
            ) from exc

        if raw in {"", "1", "r", "к", "повторить"}:
            return "retry"
        if raw in {"2", "s", "ы", "skip", "пропустить"}:
            return "skip"
        if raw in {"3", "q", "й", "stop", "стоп", "остановить"}:
            return "stop"
        print("Введите 1, 2 или 3.")


@dataclass(frozen=True)
class SourceAccount:
    row_number: int
    fio: str
    login: str
    password: str
    project_hint: str = ""
    created_date: date | None = None


@dataclass(frozen=True)
class TargetAccount:
    row_number: int
    fio: str
    login: str
    password: str
    word_reference: str
    title_override: str = ""


@dataclass(frozen=True)
class Settings:
    base_url: str
    projects_url: str
    source_accounts_file: Path
    target_accounts_file: Path
    downloads_dir: Path
    logs_dir: Path
    screenshots_dir: Path
    project_template_name: str
    headless: bool
    slow_mo_ms: int
    timeout_ms: int
    login_timeout_ms: int
    block_heavy_resources: bool
    screenshots_mode: str
    project_match_threshold: float
    stop_on_error: bool
    save_after_each_tab: bool
    strict_required_fields: bool
    interactive_account_selection: bool
    interactive_project_selection: bool


@dataclass(frozen=True)
class ProjectCandidate:
    title: str
    href: str
    score: float
    created_date: date | None = None
    created_date_text: str = ""
    status: str = ""
    order: int = 0


@dataclass
class CalendarEvent:
    task: str = ""
    title: str = ""
    deadline: str = ""
    description: str = ""
    unique_participants: str = "0"
    repeat_participants: str = "0"
    publications: str = "0"
    views: str = "0"
    additional: str = ""


@dataclass
class CofinancingEntry:
    kind: str = "partner"  # own | partner
    partner_name: str = ""
    support_type: str = ""
    expenses: str = ""
    amount: str = ""


@dataclass
class ExpenseItem:
    category: str = ""
    item: str = ""
    justification: str = ""
    price: str = ""
    quantity: str = ""


@dataclass
class ProjectData:
    source_file: Path
    title: str = ""
    region: str = ""
    scale: str = ""
    start_date: str = ""
    end_date: str = ""
    summary: str = ""
    target_groups: str = ""
    problem: str = ""
    goal: str = ""
    team_experience: str = ""
    development: str = ""
    geography: list[str] = field(default_factory=list)
    events_count: str = ""
    participants_count: str = ""
    publications_count: str = ""
    views_count: str = ""
    social_effect: str = ""
    calendar: list[CalendarEvent] = field(default_factory=list)
    cofinancing: list[CofinancingEntry] = field(default_factory=list)
    expenses: list[ExpenseItem] = field(default_factory=list)
    extra_file_descriptions: list[str] = field(default_factory=list)
    raw_fields: dict[str, list[str]] = field(default_factory=dict)
    parser_warnings: list[str] = field(default_factory=list)


@dataclass
class TransferResult:
    row_number: int
    fio: str
    operation: str
    project_name: str = ""
    word_file: str = ""
    status: str = ""
    fields_filled: int = 0
    fields_skipped: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    error: str = ""

    def skipped_text(self) -> str:
        return "; ".join(self.fields_skipped)

    def warnings_text(self) -> str:
        return "; ".join(self.warnings)


@dataclass(frozen=True)
class ParsedRow:
    section: str
    label: str
    value: str
    cells: tuple[str, ...]


FIELD_ALIASES: dict[str, tuple[str, ...]] = {
    "title": ("Название проекта",),
    "region": ("Регион реализации проекта",),
    "scale": ("Масштаб проекта", "Масштаб реализации проекта"),
    "dates": ("Дата начала и окончания проекта", "Сроки реализации проекта"),
    "start_date": ("Дата начала проекта", "Начало реализации проекта"),
    "end_date": ("Дата окончания проекта", "Окончание реализации проекта"),
    "summary": ("Краткая информация о проекте",),
    "target_groups": ("Основные целевые группы, на которые направлен проект",),
    "problem": (
        "Описание проблемы, решению/снижению которой посвящен проект",
        "Описание проблемы, решению которой посвящен проект",
    ),
    "goal": ("Основная цель проекта",),
    "team_experience": (
        "Опыт команды проекта по реализации социально значимых проектов и/или событий",
    ),
    "development": ("Перспектива развития и потенциал проекта",),
    "geography": ("География проекта",),
    "events_count": ("Количество мероприятий, проведенных в рамках проекта",),
    "participants_count": (
        "Количество участников мероприятий, вовлеченных в реализацию проекта",
        "Количество участников мероприятий",
        "Количество участников",
    ),
    "publications_count": (
        "Количество публикаций о мероприятиях проекта в средствах массовой информации, а также в информационно-телекоммуникационной сети «Интернет»",
        "Количество публикаций",
    ),
    "views_count": (
        "Количество просмотров публикаций о мероприятиях проекта в информационно-телекоммуникационной сети «Интернет»",
        "Количество просмотров публикаций",
    ),
    "social_effect": ("Социальный эффект",),
}

CALENDAR_ALIASES: dict[str, tuple[str, ...]] = {
    "task": ("Поставленная задача", "Задача"),
    "title": ("Название мероприятия",),
    "deadline": ("Крайняя дата выполнения", "Дата проведения", "Срок выполнения"),
    "description": ("Описание мероприятия",),
    "unique_participants": (
        "Количество уникальных участников",
        "Ожидаемое количество уникальных участников",
    ),
    "repeat_participants": (
        "Количество повторяющихся участников",
        "Ожидаемое количество повторяющихся участников",
    ),
    "publications": (
        "Количество публикаций",
        "Ожидаемое количество публикаций о мероприятии в СМИ и интернете",
    ),
    "views": (
        "Количество просмотров",
        "Ожидаемое количество просмотров публикаций о мероприятии в СМИ и интернете",
    ),
    "additional": ("Дополнительная информация",),
}

COFINANCING_ALIASES: dict[str, tuple[str, ...]] = {
    "partner_name": ("Название партнера", "Название партнёра", "Партнер", "Партнёр"),
    "support_type": ("Тип поддержки",),
    "expenses": ("Перечень расходов",),
    "amount": ("Сумма, руб.", "Сумма", "Сумма руб."),
}

EXPENSE_ALIASES: dict[str, tuple[str, ...]] = {
    "item": (
        "Услуга / Товар",
        "Услуга/Товар",
        "Товар / Услуга",
        "Название",
        "Наименование",
    ),
    "justification": ("Обоснование", "Описание"),
    "price": ("Цена, руб.", "Цена", "Цена руб."),
    "quantity": ("Кол-во", "Количество"),
}

# В актуальной Word-выгрузке названия категорий сокращены относительно API.
# Значения приводятся к каноническим категориям, которые затем сопоставляются
# с categoryID актуального шаблона проекта.
EXPENSE_CATEGORY_ALIASES: dict[str, tuple[str, ...]] = {
    "Расходы на создание и/или техническую поддержку сайта": (
        "Создание и/или техническая поддержка сайта",
        "Сайт и техническая поддержка",
    ),
    "Расходы на телефонную связь, мобильную связь, информационно-телекоммуникационной сети «Интернет», почтовые расходы": (
        "Связь, интернет и почтовые расходы",
        "Телефонная связь, интернет, почтовые расходы",
    ),
    "Расходы на канцелярские принадлежности": (
        "Канцелярские принадлежности",
        "Канцелярия",
    ),
    "Расходы на издательско-полиграфические услуги, в т.ч. изготовление макета, разработка дизайна": (
        "Полиграфическая продукция",
        "Издательско-полиграфические услуги",
        "Полиграфия",
    ),
    "Расходы на подарки, сувенирную продукцию": (
        "Подарки, сувенирная продукция",
        "Подарки и сувенирная продукция",
        "Сувенирная продукция",
    ),
    "Расходы на проживание и питание": (
        "Проживание и питание",
    ),
    "Транспортные расходы (приобретение авиа- и железнодорожных билетов, горюче-смазочных материалов, услуги по перевозке пассажиров)": (
        "Транспортные расходы",
        "Транспорт",
    ),
    "Расходы на аренду помещения": ("Аренда помещения",),
    "Расходы на аренду оборудования": ("Аренда оборудования",),
    "Расходы на информационные услуги (размещение информации о проекте в средствах массовой информации)": (
        "Информационные услуги",
        "Размещение информации в СМИ",
    ),
    "Расходы на закупку оборудования": (
        "Закупка оборудования",
        "Оборудование",
    ),
    "Закупка расходных материалов и оплата услуг, необходимых для реализации проекта": (
        "Расходные материалы и услуги",
        "Закупка расходных материалов",
    ),
    "Расходы на покупку и/или создание программного обеспечения": (
        "Покупка и/или создание программного обеспечения",
        "Программное обеспечение",
    ),
}


# ---------------------------------------------------------------------------
# Общие функции
# ---------------------------------------------------------------------------


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Скачивание проектов из аккаунтов myrosmol.ru и создание новых "
            "черновиков по выгруженным Word-файлам."
        )
    )
    parser.add_argument("--config", default="config_transfer.yaml")
    parser.add_argument(
        "--mode",
        choices=(MODE_DOWNLOAD, MODE_UPLOAD, MODE_ALL, MODE_PARSE),
        default="",
        help=(
            "download - скачать проекты; upload - создать черновики; "
            "all - скачать и затем перенести; parse - только проверить Word-файлы"
        ),
    )
    parser.add_argument("--row", type=int, default=0, help="Обработать одну строку XLS")
    parser.add_argument("--limit", type=int, default=0, help="Обработать первые N строк")
    parser.add_argument(
        "--auto-select",
        action="store_true",
        help=(
            "Не показывать список проектов в консоли, а автоматически выбрать "
            "проект по столбцам D-E файла исходник.xls"
        ),
    )
    parser.add_argument(
        "--all-accounts",
        action="store_true",
        help="Не спрашивать аккаунт и обработать все строки исходник.xls",
    )
    return parser.parse_args()


def choose_menu_mode() -> str:
    print("\nВыберите действие:")
    print("1 - Скачать проекты из аккаунтов файла «исходник.xls»")
    print("2 - Создать черновики по Word-файлам для аккаунтов «итог.xls»")
    print("3 - Выполнить оба этапа подряд")
    print("4 - Только проверить разбор Word-файлов без входа на сайт")
    while True:
        choice = input("Введите номер [Enter = 3]: ").strip()
        if choice == "1":
            return MODE_DOWNLOAD
        if choice == "2":
            return MODE_UPLOAD
        if choice in ("", "3"):
            return MODE_ALL
        if choice == "4":
            return MODE_PARSE
        print("Введите 1, 2, 3 или 4.")


def normalize_text(value: str) -> str:
    value = unicodedata.normalize("NFKC", value or "")
    value = value.casefold().replace("ё", "е")
    value = re.sub(r"[^0-9a-zа-я]+", " ", value, flags=re.IGNORECASE)
    return re.sub(r"\s+", " ", value).strip()


def clean_text(value: str) -> str:
    value = unicodedata.normalize("NFKC", value or "")
    value = value.replace("\r", "\n")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r" *\n *", "\n", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def safe_fs_name(value: str, max_length: int = 150) -> str:
    value = unicodedata.normalize("NFKC", value or "")
    value = re.sub(r"[\\/:*?\"<>|\r\n]+", "_", value)
    value = re.sub(r"\s+", " ", value).strip(" ._")
    return (value[:max_length] or "project").rstrip(" .")


def parse_date(value: str) -> date | None:
    value = clean_text(value)
    for pattern, fmt in (
        (r"\b\d{2}\.\d{2}\.\d{4}\b", "%d.%m.%Y"),
        (r"\b\d{4}-\d{2}-\d{2}\b", "%Y-%m-%d"),
    ):
        match = re.search(pattern, value)
        if match:
            try:
                return datetime.strptime(match.group(0), fmt).date()
            except ValueError:
                pass
    return None


def unique_path(path: Path) -> Path:
    if not path.exists():
        return path
    counter = 2
    while True:
        candidate = path.with_name(f"{path.stem} ({counter}){path.suffix}")
        if not candidate.exists():
            return candidate
        counter += 1


def resolve_relative(base: Path, value: str | Path) -> Path:
    path = Path(str(value))
    return path.resolve() if path.is_absolute() else (base / path).resolve()


def format_number(value: str, default: str = "0") -> str:
    value = clean_text(value)
    if not value:
        return default
    match = re.search(r"-?\d+(?:[.,]\d+)?", value.replace(" ", ""))
    return match.group(0).replace(",", ".") if match else default


def aliases_match(label: str, aliases: Sequence[str]) -> bool:
    normalized = normalize_text(label)
    for alias in aliases:
        target = normalize_text(alias)
        if normalized == target or normalized.startswith(target + " "):
            return True
    return False


def split_multi_value(value: str) -> list[str]:
    value = clean_text(value)
    if not value:
        return []
    parts = [clean_text(part) for part in re.split(r"\n|;", value)]
    return [part for part in parts if part]


# ---------------------------------------------------------------------------
# Конфигурация и XLS
# ---------------------------------------------------------------------------


def load_settings(config_path: Path) -> Settings:
    if not config_path.is_file():
        raise AutomationError(f"Не найден файл конфигурации: {config_path}")

    with config_path.open("r", encoding="utf-8") as stream:
        raw = yaml.safe_load(stream) or {}

    root = config_path.parent
    browser = raw.get("browser") or {}
    paths = raw.get("paths") or {}
    behavior = raw.get("behavior") or {}
    performance = raw.get("performance") or {}

    base_url = str(raw.get("base_url", DEFAULT_BASE_URL)).rstrip("/")
    projects_url = str(raw.get("projects_url", f"{base_url}/projects"))
    screenshots_mode = str(performance.get("screenshots", "errors")).casefold()
    if screenshots_mode not in {"all", "errors", "none"}:
        raise AutomationError("performance.screenshots: допустимы all, errors или none")

    return Settings(
        base_url=base_url,
        projects_url=projects_url,
        source_accounts_file=resolve_relative(
            root, raw.get("source_accounts_file", "исходник.xls")
        ),
        target_accounts_file=resolve_relative(
            root, raw.get("target_accounts_file", "итог.xls")
        ),
        downloads_dir=resolve_relative(
            root, paths.get("downloaded_projects", "скачанные_проекты")
        ),
        logs_dir=resolve_relative(root, paths.get("logs", "logs_transfer")),
        screenshots_dir=resolve_relative(
            root, paths.get("screenshots", "screenshots_transfer")
        ),
        project_template_name=str(
            raw.get("project_template_name", PROJECT_TEMPLATE_DEFAULT)
        ),
        headless=bool(browser.get("headless", False)),
        slow_mo_ms=int(browser.get("slow_mo_ms", 0)),
        timeout_ms=int(browser.get("timeout_ms", 15_000)),
        login_timeout_ms=int(browser.get("login_timeout_ms", 30_000)),
        block_heavy_resources=bool(
            performance.get("block_heavy_resources", True)
        ),
        screenshots_mode=screenshots_mode,
        project_match_threshold=float(
            behavior.get("project_match_threshold", 0.62)
        ),
        stop_on_error=bool(behavior.get("stop_on_error", False)),
        save_after_each_tab=bool(behavior.get("save_after_each_tab", False)),
        strict_required_fields=bool(
            behavior.get("strict_required_fields", False)
        ),
        interactive_account_selection=bool(
            behavior.get("interactive_account_selection", True)
        ),
        interactive_project_selection=bool(
            behavior.get("interactive_project_selection", True)
        ),
    )


def cell_to_text(cell: xlrd.sheet.Cell) -> str:
    if cell.ctype in (xlrd.XL_CELL_EMPTY, xlrd.XL_CELL_BLANK):
        return ""
    if cell.ctype == xlrd.XL_CELL_NUMBER:
        number = float(cell.value)
        if math.isfinite(number) and number.is_integer():
            return str(int(number))
        return str(number).strip()
    if cell.ctype == xlrd.XL_CELL_BOOLEAN:
        return "1" if cell.value else "0"
    return clean_text(str(cell.value))


def read_xls_rows(path: Path, max_columns: int = 6) -> list[tuple[int, list[str]]]:
    if not path.is_file():
        raise AutomationError(f"Не найден файл: {path}")
    if path.suffix.casefold() != ".xls":
        raise AutomationError(f"Файл должен быть в формате .xls: {path.name}")
    try:
        workbook = xlrd.open_workbook(path)
    except Exception as exc:
        raise AutomationError(f"Не удалось открыть {path.name}: {exc}") from exc
    if workbook.nsheets == 0:
        raise AutomationError(f"В {path.name} нет листов")

    sheet = workbook.sheet_by_index(0)
    rows: list[tuple[int, list[str]]] = []
    for row_index in range(sheet.nrows):
        values = [
            cell_to_text(sheet.cell(row_index, col)) if col < sheet.ncols else ""
            for col in range(max_columns)
        ]
        if any(values):
            rows.append((row_index + 1, values))
    return rows


def looks_like_header(values: Sequence[str]) -> bool:
    joined = " ".join(normalize_text(v) for v in values[:5])
    return "фио" in joined and "логин" in joined and "парол" in joined


def looks_like_instruction_row(values: Sequence[str]) -> bool:
    """Пропускает поясняющую строку внизу готовых шаблонов XLS."""
    if any(clean_text(value) for value in values[1:]):
        return False
    first = normalize_text(values[0] if values else "")
    return first.startswith(
        (
            "заполняйте",
            "в столбце",
            "инструкция",
            "примечание",
            "пример заполнения",
        )
    )


def read_source_accounts(path: Path) -> list[SourceAccount]:
    result: list[SourceAccount] = []
    for row_number, values in read_xls_rows(path, 5):
        if looks_like_header(values) or looks_like_instruction_row(values):
            continue
        fio, login, password, project_hint, created_text = values[:5]
        if not all((fio, login, password)):
            raise AutomationError(
                f"{path.name}, строка {row_number}: заполните A-C "
                "(ФИО, логин, пароль). Название проекта в D теперь необязательно"
            )
        result.append(
            SourceAccount(
                row_number=row_number,
                fio=fio,
                login=login,
                password=password,
                project_hint=project_hint,
                created_date=parse_date(created_text) if created_text else None,
            )
        )
    if not result:
        raise AutomationError(f"В {path.name} нет строк с аккаунтами")
    return result


def read_target_accounts(path: Path) -> list[TargetAccount]:
    result: list[TargetAccount] = []
    for row_number, values in read_xls_rows(path, 5):
        if looks_like_header(values) or looks_like_instruction_row(values):
            continue
        fio, login, password, word_reference, title_override = values[:5]
        if not all((fio, login, password, word_reference)):
            raise AutomationError(
                f"{path.name}, строка {row_number}: заполните A-D "
                "(ФИО, логин, пароль, Word-файл или название проекта)"
            )
        result.append(
            TargetAccount(
                row_number=row_number,
                fio=fio,
                login=login,
                password=password,
                word_reference=word_reference,
                title_override=title_override,
            )
        )
    if not result:
        raise AutomationError(f"В {path.name} нет строк с аккаунтами")
    return result


def filter_rows(rows: list, row_number: int, limit: int) -> list:
    if row_number:
        selected = [item for item in rows if item.row_number == row_number]
        if not selected:
            raise AutomationError(f"В XLS нет строки №{row_number}")
        return selected
    return rows[:limit] if limit > 0 else rows


def choose_source_accounts(rows: list[SourceAccount]) -> list[SourceAccount]:
    """Выбор одного или нескольких аккаунтов перед открытием браузера."""
    if len(rows) <= 1:
        return rows

    print("\nАккаунты из файла «исходник.xls»:")
    for index, account in enumerate(rows, 1):
        hint = f" — подсказка: {account.project_hint}" if account.project_hint else ""
        print(f"{index}. {account.fio} ({account.login}){hint}")
    print("A - обработать все аккаунты")

    while True:
        try:
            raw = input(
                "Введите номер аккаунта или несколько номеров через запятую: "
            ).strip()
        except EOFError as exc:
            raise AutomationError(
                "Не удалось прочитать выбор аккаунта из консоли"
            ) from exc

        if raw.casefold() in {"a", "а", "all", "все"}:
            return rows

        selected: list[SourceAccount] = []
        invalid = False
        seen: set[int] = set()
        for token in re.split(r"[,;\s]+", raw):
            if not token:
                continue
            if not token.isdigit():
                invalid = True
                break
            index = int(token)
            if not 1 <= index <= len(rows):
                invalid = True
                break
            if index not in seen:
                seen.add(index)
                selected.append(rows[index - 1])
        if selected and not invalid:
            return selected
        print(f"Введите номер от 1 до {len(rows)} или A для всех аккаунтов.")


# ---------------------------------------------------------------------------
# Разбор Word-файла
# ---------------------------------------------------------------------------


def is_placeholder_value(value: str) -> bool:
    normalized = normalize_text(value)
    if not normalized:
        return True
    prefixes = (
        "обязательное поле",
        "необязательное поле",
        "автоматическое поле",
        "отображается автоматически",
        "переносится автоматически",
        "на данном этапе",
        "добавить поле",
        "добавить задачу",
        "добавить мероприятие",
        "добавить строку",
        "добавить файл",
    )
    return any(normalized.startswith(prefix) for prefix in prefixes)


def detect_section(text: str) -> str | None:
    match = re.search(r"Вкладка\s*[«\"]([^»\"]+)[»\"]", text, re.I)
    if match:
        return clean_text(match.group(1))
    return None


def all_known_labels() -> list[str]:
    labels: list[str] = []
    for mapping in (
        FIELD_ALIASES,
        CALENDAR_ALIASES,
        COFINANCING_ALIASES,
        EXPENSE_ALIASES,
    ):
        for aliases in mapping.values():
            labels.extend(aliases)
    labels.extend(
        [
            "Собственные средства",
            "Партнер",
            "Партнёр",
            "Описание файла",
            *EXPENSE_CATEGORIES,
        ]
    )
    return sorted(set(labels), key=len, reverse=True)


KNOWN_LABELS = all_known_labels()


def split_label_value_block(text: str) -> tuple[str, str] | None:
    text = clean_text(text)
    normalized = normalize_text(text)
    for label in KNOWN_LABELS:
        normalized_label = normalize_text(label)
        if normalized == normalized_label:
            return label, ""
        if normalized.startswith(normalized_label + " "):
            # Сохраняем исходный текст после длины видимого ярлыка, насколько возможно.
            pattern = re.compile(r"^\s*" + re.escape(label) + r"\s*[:\-–—]?\s*", re.I)
            match = pattern.match(text)
            if match:
                return label, clean_text(text[match.end() :])
    return None


def choose_row_value(label: str, cells: Sequence[str]) -> str:
    normalized_label = normalize_text(label)
    candidates: list[str] = []
    for cell in cells[1:]:
        cell = clean_text(cell)
        if not cell or normalize_text(cell) == normalized_label:
            continue
        if detect_section(cell):
            continue
        if is_placeholder_value(cell):
            continue
        candidates.append(cell)
    return candidates[-1] if candidates else ""


def _paragraph_label_value(text: str) -> tuple[str, str] | None:
    """Разделяет обычный абзац вида «Подпись: значение»."""
    text = clean_text(text)
    match = re.match(r"^(.{1,260}?):\s*(.*)$", text, re.S)
    if not match:
        return None
    return clean_text(match.group(1)), clean_text(match.group(2))


def _is_paragraph_service_line(text: str) -> bool:
    normalized = normalize_text(text)
    if not normalized:
        return True
    if re.match(r'^блок\s+[«"].+[»"]$', text, re.I):
        return True
    if re.match(r'^запись\s*№\s*\d+', text, re.I):
        return True
    if re.match(r'^тип\s+[«"].+[»"]$', text, re.I):
        return True
    return normalized in {
        "поле",
        "добавить",
        "добавить задачу",
        "добавить мероприятие",
        "добавить файл",
        "партнера",
        "партнер",
        "команда",
        "наставники",
    }


def _result_label_from_block(block: str) -> str | None:
    for key in (
        "events_count",
        "participants_count",
        "publications_count",
        "views_count",
    ):
        aliases = FIELD_ALIASES[key]
        if aliases_match(block, aliases):
            return aliases[0]
    return None


def extract_rows_from_docx(path: Path) -> list[ParsedRow]:
    try:
        document = Document(path)
    except Exception as exc:
        raise AutomationError(f"Не удалось открыть Word-файл {path.name}: {exc}") from exc

    parsed: list[ParsedRow] = []
    section = ""

    # Старые выгрузки могут содержать таблицы — их обработка сохраняется.
    for table in document.tables:
        for row in table.rows:
            raw_cells = [clean_text(cell.text) for cell in row.cells]
            cells: list[str] = []
            for cell in raw_cells:
                if not cells or normalize_text(cell) != normalize_text(cells[-1]):
                    cells.append(cell)
            if not any(cells):
                continue

            section_candidate = next(
                (detect_section(cell) for cell in cells if detect_section(cell)), None
            )
            if section_candidate:
                section = section_candidate
                continue

            label = cells[0]
            value = choose_row_value(label, cells)
            if len(cells) == 1:
                split = split_label_value_block(cells[0])
                if split:
                    label, value = split

            parsed.append(
                ParsedRow(
                    section=section,
                    label=clean_text(label),
                    value=clean_text(value),
                    cells=tuple(cells),
                )
            )

    # Актуальная выгрузка myrosmol.ru представляет заявку последовательностью
    # обычных абзацев: подпись поля и значение чаще всего находятся в разных
    # абзацах. Поэтому здесь используется контекстный разбор с запоминанием
    # текущей вкладки, блока и ожидающего значения поля.
    paragraph_section = ""
    current_block = ""
    pending_label = ""
    pending_result_label = ""
    geography_region = ""

    def append_row(label: str, value: str, source_text: str) -> None:
        value = clean_text(value)
        if not value or is_placeholder_value(value):
            return
        parsed.append(
            ParsedRow(
                section=paragraph_section,
                label=clean_text(label),
                value=value,
                cells=(source_text,),
            )
        )

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if not text:
            continue

        section_candidate = detect_section(text)
        if section_candidate:
            paragraph_section = section_candidate
            current_block = ""
            pending_label = ""
            pending_result_label = ""
            geography_region = ""
            continue

        block_match = re.match(r'^Блок\s+[«"](.+?)[»"]$', text, re.I)
        if block_match:
            current_block = clean_text(block_match.group(1))
            pending_label = ""
            pending_result_label = ""
            if normalize_text(current_block) != normalize_text("География проекта"):
                geography_region = ""
            if (
                normalize_text(paragraph_section) == normalize_text("Софинансирование")
                and normalize_text(current_block) in {
                    normalize_text("Собственные средства"),
                    normalize_text("Партнер"),
                    normalize_text("Партнёр"),
                }
            ):
                parsed.append(
                    ParsedRow(
                        section=paragraph_section,
                        label=current_block,
                        value="",
                        cells=(text,),
                    )
                )
            continue

        category_match = re.match(r'^Категория\s+[«"](.+?)[»"]$', text, re.I)
        if category_match and normalize_text(paragraph_section) == normalize_text("Расходы"):
            category = best_expense_category(category_match.group(1)) or clean_text(
                category_match.group(1)
            )
            parsed.append(
                ParsedRow(
                    section=paragraph_section,
                    label=category,
                    value="",
                    cells=(text,),
                )
            )
            pending_label = ""
            continue

        if re.match(r'^Запись\s*№\s*\d+', text, re.I):
            pending_label = ""
            continue
        if re.match(r'^Тип\s+[«"].+[»"]$', text, re.I):
            continue

        # Сначала определяем, является ли текущий абзац новой известной
        # подписью. Если нет, а предыдущее поле ждёт значение, весь текущий
        # абзац считается значением даже при наличии двоеточий внутри текста.
        split = split_label_value_block(text)

        if pending_result_label:
            if split is None:
                append_row(pending_result_label, text, text)
                pending_result_label = ""
                continue
            pending_result_label = ""

        if pending_label:
            if split is None:
                if pending_label == "__geography_address__":
                    location = ", ".join(
                        part for part in (geography_region, text) if clean_text(part)
                    )
                    append_row("География проекта", location or text, text)
                elif not _is_paragraph_service_line(pending_label):
                    append_row(pending_label, text, text)
                pending_label = ""
                continue
            # Новая известная подпись означает, что предыдущее поле было пустым.
            pending_label = ""

        # Известные длинные подписи обрабатываются раньше универсального
        # разделения по двоеточию.
        if split:
            label, value = split
            if value:
                append_row(label, value, text)
                pending_label = ""
            else:
                pending_label = label
            continue

        generic = _paragraph_label_value(text)
        if generic:
            raw_label, value = generic
            label_norm = normalize_text(raw_label)
            if label_norm in {"http", "https"}:
                continue

            # В блоках результатов одно и то же поле «Плановое количество»
            # относится к разным показателям. Преобразуем его в название блока.
            if (
                normalize_text(paragraph_section) == normalize_text("Результаты")
                and label_norm == normalize_text("Плановое количество")
            ):
                pending_result_label = _result_label_from_block(current_block) or ""
                if value and pending_result_label:
                    append_row(pending_result_label, value, text)
                    pending_result_label = ""
                pending_label = ""
                continue

            if _is_paragraph_service_line(raw_label):
                pending_label = ""
                continue

            # География хранится внутри составного блока. Формируем одно
            # значение «регион, адрес», пригодное для API-сопоставления.
            if (
                normalize_text(paragraph_section) == normalize_text("О проекте")
                and normalize_text(current_block) == normalize_text("География проекта")
            ):
                if label_norm.startswith(normalize_text("Выберите регион или федеральный округ")):
                    geography_region = value
                    if value:
                        append_row("География проекта", value, text)
                    pending_label = ""
                    continue
                if label_norm == normalize_text("Адрес"):
                    if value:
                        location = ", ".join(
                            part for part in (geography_region, value) if clean_text(part)
                        )
                        append_row("География проекта", location or value, text)
                    else:
                        pending_label = "__geography_address__"
                    continue

            if value:
                append_row(raw_label, value, text)
                pending_label = ""
            else:
                pending_label = raw_label
            continue

        # Значение, вынесенное в следующий абзац после подписи.
        if pending_result_label:
            append_row(pending_result_label, text, text)
            pending_result_label = ""
            continue

        if pending_label:
            if pending_label == "__geography_address__":
                location = ", ".join(
                    part for part in (geography_region, text) if clean_text(part)
                )
                append_row("География проекта", location or text, text)
            elif not _is_paragraph_service_line(pending_label):
                append_row(pending_label, text, text)
            pending_label = ""
            continue

    return parsed


def raw_add(project: ProjectData, label: str, value: str) -> None:
    if value:
        project.raw_fields.setdefault(normalize_text(label), []).append(value)


def first_field(rows: Sequence[ParsedRow], aliases: Sequence[str]) -> str:
    values = [row.value for row in rows if aliases_match(row.label, aliases) and row.value]
    return values[0] if values else ""


def rows_for_section(rows: Sequence[ParsedRow], *section_names: str) -> list[ParsedRow]:
    allowed = {normalize_text(name) for name in section_names}
    return [row for row in rows if normalize_text(row.section) in allowed]


def parse_date_range(value: str) -> tuple[str, str]:
    value = clean_text(value)
    if not value:
        return "", ""

    tokens = re.findall(
        r"(?:\d{2}\.\d{2}\.\d{4}|\d{2}\.\d{4}|\d{4}-\d{2}-\d{2}|\d{4}-\d{2})",
        value,
    )
    if len(tokens) >= 2:
        return tokens[0], tokens[1]
    if len(tokens) == 1:
        return tokens[0], tokens[0]

    parts = [clean_text(part) for part in re.split(r"\s+[–—-]\s+|\n", value)]
    parts = [part for part in parts if part]
    if len(parts) >= 2:
        return parts[0], parts[1]
    return value, value


def match_key(label: str, aliases_map: dict[str, tuple[str, ...]]) -> str | None:
    for key, aliases in aliases_map.items():
        if aliases_match(label, aliases):
            return key
    return None


def parse_calendar(rows: Sequence[ParsedRow]) -> list[CalendarEvent]:
    events: list[CalendarEvent] = []
    current_task = ""
    current: CalendarEvent | None = None

    def finish() -> None:
        nonlocal current
        if current and (current.title or current.description):
            if not current.task:
                current.task = current_task
            events.append(current)
        current = None

    for row in rows:
        key = match_key(row.label, CALENDAR_ALIASES)
        if key is None or not row.value:
            continue
        if key == "task":
            if current and current.title:
                finish()
            current_task = row.value
            if current is None:
                current = CalendarEvent(task=current_task)
            else:
                current.task = current_task
            continue
        if key == "title":
            if current and current.title:
                finish()
            current = CalendarEvent(task=current_task, title=row.value)
            continue
        if current is None:
            current = CalendarEvent(task=current_task)
        setattr(current, key, row.value)

    finish()
    return events


def parse_cofinancing(rows: Sequence[ParsedRow]) -> list[CofinancingEntry]:
    entries: list[CofinancingEntry] = []
    mode = ""
    current: CofinancingEntry | None = None

    def finish() -> None:
        nonlocal current
        if current and (current.expenses or current.amount or current.partner_name):
            entries.append(current)
        current = None

    for row in rows:
        label_norm = normalize_text(row.label)
        if label_norm == normalize_text("Собственные средства"):
            finish()
            mode = "own"
            current = CofinancingEntry(kind="own")
            continue
        if label_norm in {
            normalize_text("Партнер"),
            normalize_text("Партнёр"),
            normalize_text("Партнеры"),
        } and not row.value:
            finish()
            mode = "partner"
            current = None
            continue

        key = match_key(row.label, COFINANCING_ALIASES)
        if key is None or not row.value:
            continue

        if mode == "own":
            if current is None or current.kind != "own":
                current = CofinancingEntry(kind="own")
            if key in {"expenses", "amount"}:
                setattr(current, key, row.value)
            continue

        mode = "partner"
        if key == "support_type":
            # В актуальной выгрузке «Тип поддержки» открывает новую запись и
            # расположен перед названием партнёра.
            if current and any((current.partner_name, current.expenses, current.amount)):
                finish()
            if current is None:
                current = CofinancingEntry(kind="partner")
            current.support_type = row.value
            continue

        if key == "partner_name":
            if current is None:
                current = CofinancingEntry(kind="partner")
            elif current.partner_name or current.expenses or current.amount:
                finish()
                current = CofinancingEntry(kind="partner")
            current.partner_name = row.value
            continue

        if current is None:
            current = CofinancingEntry(kind="partner")
        setattr(current, key, row.value)

    finish()
    return entries


def best_expense_category(text: str) -> str | None:
    normalized = normalize_text(text)
    if not normalized:
        return None

    exact: dict[str, str] = {}
    for canonical in EXPENSE_CATEGORIES:
        exact[normalize_text(canonical)] = canonical
        for alias in EXPENSE_CATEGORY_ALIASES.get(canonical, ()):
            exact[normalize_text(alias)] = canonical

    if normalized in exact:
        return exact[normalized]
    for key, original in exact.items():
        if key and len(key) >= 8 and (key in normalized or normalized in key):
            return original
    return None


def parse_expenses(rows: Sequence[ParsedRow]) -> list[ExpenseItem]:
    items: list[ExpenseItem] = []
    current_category = ""
    current: ExpenseItem | None = None

    def finish() -> None:
        nonlocal current
        if current and current.item:
            items.append(current)
        current = None

    for row in rows:
        # Категория определяется только по самой строке категории. Нельзя
        # искать её в описании товара: слова «транспортировка», «оборудование»
        # и т.п. внутри обоснования ранее ошибочно переключали категорию.
        category = best_expense_category(row.label)
        if category is None and "категор" in normalize_text(row.label):
            category = best_expense_category(row.value)

        if category:
            finish()
            current_category = category
            if match_key(row.label, EXPENSE_ALIASES) is None:
                continue

        # Поддержка старых табличных выгрузок, где выбранная категория могла
        # находиться в отдельной ячейке строки.
        if not category and not match_key(row.label, EXPENSE_ALIASES):
            for cell in row.cells:
                exact_category = best_expense_category(cell)
                if exact_category:
                    finish()
                    current_category = exact_category
                    break

        key = match_key(row.label, EXPENSE_ALIASES)
        if key is None or not row.value:
            continue
        if key == "item":
            finish()
            current = ExpenseItem(category=current_category, item=row.value)
            continue
        if current is None:
            current = ExpenseItem(category=current_category)
        setattr(current, key, row.value)

    finish()
    return items


def parse_project_docx(path: Path, title_override: str = "") -> ProjectData:
    rows = extract_rows_from_docx(path)
    project = ProjectData(source_file=path)
    for row in rows:
        raw_add(project, row.label, row.value)

    general_rows = rows_for_section(rows, "Общее")
    about_rows = rows_for_section(rows, "О проекте")
    results_rows = rows_for_section(rows, "Результаты")
    calendar_rows = rows_for_section(rows, "Календарный план")
    cofinancing_rows = rows_for_section(rows, "Софинансирование")
    expense_rows = rows_for_section(rows, "Расходы")
    extra_rows = rows_for_section(rows, "Доп. файлы", "Дополнительные файлы")

    project.title = title_override or (
        first_field(general_rows, FIELD_ALIASES["title"])
        or first_field(rows, FIELD_ALIASES["title"])
    )
    project.region = (
        first_field(general_rows, FIELD_ALIASES["region"])
        or first_field(rows, FIELD_ALIASES["region"])
    )
    project.scale = (
        first_field(general_rows, FIELD_ALIASES["scale"])
        or first_field(rows, FIELD_ALIASES["scale"])
    )

    date_range = first_field(general_rows or rows, FIELD_ALIASES["dates"])
    project.start_date, project.end_date = parse_date_range(date_range)
    project.start_date = (
        first_field(general_rows or rows, FIELD_ALIASES["start_date"]) or project.start_date
    )
    project.end_date = (
        first_field(general_rows or rows, FIELD_ALIASES["end_date"]) or project.end_date
    )

    project.summary = first_field(about_rows or rows, FIELD_ALIASES["summary"])
    project.target_groups = first_field(about_rows or rows, FIELD_ALIASES["target_groups"])
    project.problem = first_field(about_rows or rows, FIELD_ALIASES["problem"])
    project.goal = first_field(about_rows or rows, FIELD_ALIASES["goal"])
    project.team_experience = first_field(about_rows or rows, FIELD_ALIASES["team_experience"])
    project.development = first_field(about_rows or rows, FIELD_ALIASES["development"])
    geography_values = [
        row.value
        for row in (about_rows or rows)
        if aliases_match(row.label, FIELD_ALIASES["geography"]) and row.value
    ]
    # The paragraph export may first contain only the region and then the full
    # «region, address» value. Keep unique values and prefer detailed addresses.
    unique_geography: list[str] = []
    for value in geography_values:
        for part in split_multi_value(value):
            normalized = normalize_text(part)
            if normalized and all(normalize_text(existing) != normalized for existing in unique_geography):
                unique_geography.append(part)
    detailed = [value for value in unique_geography if "," in value or "г." in value.casefold()]
    project.geography = detailed or unique_geography
    project.events_count = first_field(results_rows or rows, FIELD_ALIASES["events_count"])
    project.participants_count = first_field(
        results_rows or rows, FIELD_ALIASES["participants_count"]
    )
    project.publications_count = first_field(
        results_rows or rows, FIELD_ALIASES["publications_count"]
    )
    project.views_count = first_field(results_rows or rows, FIELD_ALIASES["views_count"])
    project.social_effect = first_field(results_rows or rows, FIELD_ALIASES["social_effect"])

    project.calendar = parse_calendar(calendar_rows)
    project.cofinancing = parse_cofinancing(cofinancing_rows)
    project.expenses = parse_expenses(expense_rows)
    project.extra_file_descriptions = [
        row.value
        for row in extra_rows
        if aliases_match(row.label, ("Описание файла",)) and row.value
    ]

    if not project.title:
        project.title = path.stem
        project.parser_warnings.append(
            "Название проекта не найдено в Word; использовано имя файла"
        )
    if not project.calendar:
        project.parser_warnings.append("В Word не распознан календарный план")
    if not project.expenses:
        project.parser_warnings.append("В Word не распознаны строки сметы")

    return project


def resolve_word_file(reference: str, settings: Settings) -> Path:
    direct = Path(reference)
    candidates: list[Path] = []
    if direct.is_absolute():
        candidates.append(direct)
    else:
        candidates.extend(
            [
                settings.target_accounts_file.parent / direct,
                settings.downloads_dir / direct,
            ]
        )
    if not direct.suffix:
        candidates.extend(path.with_suffix(".docx") for path in list(candidates))
        candidates.extend(path.with_suffix(".doc") for path in list(candidates))

    for candidate in candidates:
        if candidate.is_file():
            return candidate.resolve()

    if not settings.downloads_dir.is_dir():
        raise AutomationError(
            f"Не найдена папка скачанных проектов: {settings.downloads_dir}"
        )

    normalized_reference = normalize_text(Path(reference).stem)
    matches = [
        path
        for path in settings.downloads_dir.iterdir()
        if path.is_file()
        and path.suffix.casefold() in {".docx", ".doc"}
        and (
            normalize_text(path.stem) == normalized_reference
            or normalized_reference in normalize_text(path.stem)
            or normalize_text(path.stem) in normalized_reference
        )
    ]
    if len(matches) == 1:
        return matches[0].resolve()
    if len(matches) > 1:
        raise AutomationError(
            f"Для «{reference}» найдено несколько Word-файлов: "
            + ", ".join(path.name for path in matches)
        )
    raise AutomationError(f"Word-файл «{reference}» не найден")


# ---------------------------------------------------------------------------
# Playwright: общие операции
# ---------------------------------------------------------------------------


async def wait_visible(locator: Locator, timeout_ms: int = 2500) -> bool:
    try:
        await locator.first.wait_for(state="visible", timeout=timeout_ms)
        return True
    except PlaywrightTimeoutError:
        return False


async def first_visible(
    locators: Iterable[Locator], timeout_ms: int = 2500
) -> Locator | None:
    items = [locator.first for locator in locators]
    deadline = asyncio.get_running_loop().time() + max(timeout_ms, 0) / 1000
    while True:
        for locator in items:
            try:
                if await locator.is_visible(timeout=0):
                    return locator
            except Exception:
                continue
        if asyncio.get_running_loop().time() >= deadline:
            return None
        await asyncio.sleep(0.05)


async def take_screenshot(
    page: Page,
    settings: Settings,
    row_number: int,
    fio: str,
    name: str,
) -> None:
    mode = settings.screenshots_mode
    is_error = "ошиб" in normalize_text(name)
    if mode == "none" or (mode == "errors" and not is_error):
        return
    directory = settings.screenshots_dir / f"{row_number:04d}_{safe_fs_name(fio)}"
    directory.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%H%M%S")
    try:
        await page.screenshot(
            path=str(directory / f"{stamp}_{safe_fs_name(name)}.png"),
            full_page=False,
        )
    except Exception:
        pass


async def locate_login_input(page: Page, timeout_ms: int = 2500) -> Locator | None:
    locator = page.locator(
        "input[autocomplete='username'], input[type='email'], input[type='tel'], "
        "input[name*='login' i], input[name*='email' i], input[name*='phone' i]"
    )
    if await wait_visible(locator, timeout_ms):
        return locator.first
    fallback = page.locator("input[type='text']")
    if await wait_visible(fallback, 700):
        return fallback.first
    return None


async def locate_password_input(page: Page, timeout_ms: int = 2500) -> Locator | None:
    locator = page.locator(
        "input[type='password'], input[autocomplete='current-password']"
    )
    if await wait_visible(locator, timeout_ms):
        return locator.first
    return None


async def submit_login(page: Page) -> bool:
    submit = page.locator("button[type='submit'], input[type='submit']")
    if await wait_visible(submit, 1200):
        await submit.first.click()
        return True
    target = await first_visible(
        [
            page.get_by_role("button", name=re.compile(r"Войти|Продолжить|Далее", re.I)),
            page.get_by_role("link", name=re.compile(r"Войти|Продолжить|Далее", re.I)),
        ],
        1500,
    )
    if target:
        await target.click()
        return True
    return False


async def wait_login_completed(page: Page, settings: Settings) -> None:
    deadline = asyncio.get_running_loop().time() + settings.login_timeout_ms / 1000
    while asyncio.get_running_loop().time() < deadline:
        password = page.locator(
            "input[type='password'], input[autocomplete='current-password']"
        )
        try:
            password_visible = await password.first.is_visible(timeout=200)
        except Exception:
            password_visible = False
        auth_url = any(
            token in page.url.casefold()
            for token in ("/login", "/auth", "signin", "sign-in")
        )
        if not password_visible and not auth_url:
            return
        await page.wait_for_timeout(200)
    raise AutomationError(
        "Авторизация не завершена: возможны неверные данные, CAPTCHA, код "
        "подтверждения или изменение формы входа"
    )


async def login_account(
    page: Page, login: str, password: str, settings: Settings
) -> None:
    await page.goto(
        settings.base_url,
        wait_until="domcontentloaded",
        timeout=settings.timeout_ms,
    )
    login_input = await locate_login_input(page, 700)
    password_input = await locate_password_input(page, 700)

    if login_input is None or password_input is None:
        login_button = await first_visible(
            [
                page.get_by_role("button", name=LOGIN_BUTTON_PATTERN),
                page.get_by_role("link", name=LOGIN_BUTTON_PATTERN),
            ],
            1600,
        )
        if login_button:
            await login_button.click()

        password_input = await locate_password_input(page, 1800)
        if password_input is None:
            mode = await first_visible(
                [
                    page.get_by_role("button", name=PASSWORD_MODE_PATTERN),
                    page.get_by_role("link", name=PASSWORD_MODE_PATTERN),
                    page.get_by_text(PASSWORD_MODE_PATTERN),
                ],
                1500,
            )
            if mode:
                await mode.click()

        login_input = await locate_login_input(page, 2500)
        password_input = await locate_password_input(page, 2500)

    if login_input is None or password_input is None:
        raise AutomationError("Не найдена стандартная форма входа")

    await login_input.fill(login)
    await password_input.fill(password)
    if not await submit_login(page):
        raise AutomationError("Не найдена кнопка отправки формы входа")
    await wait_login_completed(page, settings)
    await page.reload(wait_until="domcontentloaded", timeout=settings.timeout_ms)


async def open_projects(page: Page, settings: Settings, require_cards: bool) -> None:
    await page.goto(
        settings.projects_url,
        wait_until="domcontentloaded",
        timeout=settings.timeout_ms,
    )
    if any(token in page.url.casefold() for token in ("/login", "/auth", "signin")):
        raise AutomationError("После входа сайт снова открыл страницу авторизации")

    if require_cards:
        cards = page.locator("div.base-card__details")
        try:
            await cards.first.wait_for(
                state="visible", timeout=min(settings.timeout_ms, 7000)
            )
        except PlaywrightTimeoutError as exc:
            raise AutomationError("Карточки проектов не загрузились") from exc


async def load_all_project_cards(page: Page) -> None:
    """Догружает карточки при пагинации или бесконечной прокрутке."""
    cards = page.locator("div.base-card__details")
    last_count = -1
    stable_rounds = 0
    more_pattern = re.compile(
        r"Показать ещё|Показать еще|Загрузить ещё|Загрузить еще|"
        r"Показать больше|Ещё проекты|Еще проекты",
        re.I,
    )

    for _ in range(40):
        current_count = await cards.count()
        more_button = await first_visible(
            [
                page.get_by_role("button", name=more_pattern),
                page.get_by_role("link", name=more_pattern),
                page.get_by_text(more_pattern),
            ],
            250,
        )
        clicked_more = False
        if more_button is not None:
            try:
                await more_button.scroll_into_view_if_needed()
                await more_button.click()
                clicked_more = True
            except Exception:
                clicked_more = False

        try:
            await page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
        except Exception:
            pass
        await page.wait_for_timeout(350 if clicked_more else 250)
        new_count = await cards.count()

        if new_count == last_count == current_count and not clicked_more:
            stable_rounds += 1
        else:
            stable_rounds = 0
        last_count = new_count
        if stable_rounds >= 3:
            break

    try:
        await page.evaluate("window.scrollTo(0, 0)")
    except Exception:
        pass


def project_score(query: str, title: str) -> float:
    q = normalize_text(query)
    t = normalize_text(title)
    if not q or not t:
        return 0.0
    if q == t:
        return 1.0
    if q in t:
        return 0.94 + min(len(q) / max(len(t), 1), 1.0) * 0.05
    if t in q:
        return 0.90 + min(len(t) / max(len(q), 1), 1.0) * 0.05
    ratio = SequenceMatcher(None, q, t).ratio()
    q_tokens = set(q.split())
    t_tokens = set(t.split())
    overlap = len(q_tokens & t_tokens) / max(len(q_tokens | t_tokens), 1)
    return ratio * 0.65 + overlap * 0.35


def clean_project_title(value: str) -> str:
    value = re.sub(r"\s+", " ", value or "").strip()
    value = re.sub(
        r"\s+(Черновик|Подан|Подана|Архив|Архивирован|Победитель|Не победитель)\s*$",
        "",
        value,
        flags=re.I,
    )
    return value.strip()


async def collect_project_candidates(
    page: Page, project_hint: str = ""
) -> list[ProjectCandidate]:
    anchors = page.locator(
        "div.base-card__details a[href^='/projects/'], "
        "div.base-card__details a[href*='myrosmol.ru/projects/']"
    )
    try:
        raw_items = await anchors.evaluate_all(
            r"""
            anchors => anchors.slice(0, 2000).map((anchor, order) => {
                const card = anchor.closest('div.base-card__details');
                const heading = anchor.querySelector('h3');
                let title = '';
                if (heading) {
                    title = Array.from(heading.childNodes)
                        .filter(node => node.nodeType === Node.TEXT_NODE)
                        .map(node => node.textContent || '').join(' ');
                }
                if (!title.trim()) title = heading?.innerText || anchor.innerText || '';
                const text = card?.innerText || '';
                const dateMatch = text.match(/\b\d{2}\.\d{2}\.\d{4}\b/);
                const statusMatch = text.match(
                    /(?:^|\n)\s*(Черновик|Подан(?:а)?|Архив(?:ирован)?|Победитель|Не победитель)\s*(?:$|\n)/i
                );
                return {
                    href: anchor.getAttribute('href') || '',
                    title,
                    dateText: dateMatch ? dateMatch[0] : '',
                    status: statusMatch ? statusMatch[1] : '',
                    order
                };
            })
            """
        )
    except Exception as exc:
        raise AutomationError(f"Не удалось прочитать карточки проектов: {exc}") from exc

    by_href: dict[str, ProjectCandidate] = {}
    for item in raw_items:
        href = str(item.get("href", "")).strip()
        title = clean_project_title(str(item.get("title", "")))
        if not href or len(title) < 2:
            continue
        parsed_date = parse_date(str(item.get("dateText", "")))
        candidate = ProjectCandidate(
            title=title,
            href=urljoin(page.url, href),
            score=project_score(project_hint, title) if project_hint else 0.0,
            created_date=parsed_date,
            created_date_text=(
                parsed_date.strftime("%d.%m.%Y") if parsed_date else ""
            ),
            status=clean_text(str(item.get("status", ""))),
            order=int(item.get("order", 0)),
        )
        previous = by_href.get(candidate.href)
        if previous is None or candidate.score > previous.score:
            by_href[candidate.href] = candidate

    candidates = list(by_href.values())
    if project_hint:
        return sorted(
            candidates,
            key=lambda item: (item.score, item.created_date or date.min),
            reverse=True,
        )
    return sorted(candidates, key=lambda item: item.order)


async def choose_candidate_automatically(
    account: SourceAccount,
    candidates: list[ProjectCandidate],
    settings: Settings,
) -> ProjectCandidate:
    eligible = [
        item for item in candidates if item.score >= settings.project_match_threshold
    ]
    if not eligible:
        raise AutomationError(
            f"Не найден проект, совпадающий с «{account.project_hint}»"
        )

    top = eligible[0]
    same_title = [
        item
        for item in eligible
        if normalize_text(item.title) == normalize_text(top.title)
    ]
    if len(same_title) == 1:
        if len(eligible) > 1:
            second = eligible[1]
            if abs(top.score - second.score) < 0.015 and normalize_text(
                top.title
            ) != normalize_text(second.title):
                raise AutomationError(
                    "Найдены два почти одинаковых названия: "
                    f"«{top.title}» и «{second.title}»"
                )
        return top

    if account.created_date:
        matches = [
            item for item in same_title if item.created_date == account.created_date
        ]
        if len(matches) == 1:
            return matches[0]
        available = ", ".join(
            item.created_date_text or "дата не определена" for item in same_title
        )
        raise AutomationError(
            f"Для даты {account.created_date.strftime('%d.%m.%Y')} проект не найден. "
            f"Доступные даты: {available}"
        )

    if any(item.created_date is None for item in same_title):
        raise AutomationError(
            "Найдено несколько проектов с одинаковым названием, но дата создания "
            "определена не у всех. Заполните столбец E в исходник.xls"
        )

    print(
        f"\nДля {account.fio} найдено несколько проектов «{top.title}»:"
    )
    ordered = sorted(
        same_title, key=lambda item: item.created_date or date.min, reverse=True
    )
    for index, item in enumerate(ordered, 1):
        print(f"{index}. {item.created_date_text}")
    while True:
        raw = (
            await asyncio.to_thread(
                input,
                "Введите дату создания ДД.ММ.ГГГГ (Q - пропустить): ",
            )
        ).strip()
        if raw.casefold() in {"q", "й", "пропустить"}:
            raise AutomationError("Строка пропущена пользователем")
        selected = parse_date(raw)
        matches = [item for item in ordered if item.created_date == selected]
        if len(matches) == 1:
            return matches[0]
        print("Проект с такой датой не найден")


def project_candidate_line(index: int, candidate: ProjectCandidate) -> str:
    details = []
    if candidate.created_date_text:
        details.append(candidate.created_date_text)
    if candidate.status:
        details.append(candidate.status)
    suffix = f" ({', '.join(details)})" if details else ""
    return f"{index}. {candidate.title}{suffix}"


async def choose_candidate_in_console(
    account: SourceAccount, candidates: list[ProjectCandidate]
) -> ProjectCandidate:
    if not candidates:
        raise AutomationError("В аккаунте не найдено ни одного проекта")

    print(f"\nПроекты аккаунта {account.fio}:")
    for index, candidate in enumerate(candidates, 1):
        print(project_candidate_line(index, candidate))
    if account.project_hint:
        print(f"Подсказка из столбца D: {account.project_hint}")

    while True:
        try:
            raw = (
                await asyncio.to_thread(
                    input,
                    "Введите номер проекта или часть названия (Q - пропустить): ",
                )
            ).strip()
        except EOFError as exc:
            raise AutomationError(
                "Не удалось прочитать выбор проекта из консоли"
            ) from exc

        if raw.casefold() in {"q", "й", "пропустить"}:
            raise AutomationError("Строка пропущена пользователем")

        if raw.isdigit():
            index = int(raw)
            if 1 <= index <= len(candidates):
                return candidates[index - 1]
            print(f"Введите номер от 1 до {len(candidates)}.")
            continue

        query = normalize_text(raw)
        if not query:
            print("Введите номер или часть названия проекта.")
            continue

        matches = [
            (index, candidate)
            for index, candidate in enumerate(candidates, 1)
            if query in normalize_text(candidate.title)
        ]
        if len(matches) == 1:
            return matches[0][1]
        if len(matches) > 1:
            print("Найдено несколько проектов:")
            for index, candidate in matches:
                print(project_candidate_line(index, candidate))
            print("Укажите номер нужного проекта из общего списка.")
            continue

        ranked = sorted(
            candidates,
            key=lambda item: project_score(raw, item.title),
            reverse=True,
        )[:5]
        print("Точного совпадения нет. Ближайшие варианты:")
        for candidate in ranked:
            index = candidates.index(candidate) + 1
            print(project_candidate_line(index, candidate))


async def open_candidate_project(page: Page, candidate: ProjectCandidate) -> None:
    path = urlparse(candidate.href).path
    locator = page.locator(
        f"div.base-card__details a[href='{path}'], "
        f"div.base-card__details a[href='{candidate.href}']"
    )
    if await locator.count() == 0:
        raise AutomationError("Карточка выбранного проекта исчезла со страницы")
    link = locator.first
    await link.scroll_into_view_if_needed()
    await link.click()
    try:
        await page.wait_for_load_state("domcontentloaded", timeout=5000)
    except PlaywrightTimeoutError:
        pass


async def dismiss_cookie_banner(page: Page) -> None:
    """Закрывает баннер cookie, если он перекрывает элементы формы."""
    accept = page.get_by_role(
        "button", name=re.compile(r"^\s*Принять\s*$", re.I)
    )
    if not await wait_visible(accept, 500):
        return
    try:
        await accept.first.click(timeout=1500)
    except Exception:
        try:
            await accept.first.evaluate("element => element.click()")
        except Exception:
            pass
    await page.wait_for_timeout(100)


async def find_enabled_form_action(
    page: Page,
    pattern: re.Pattern[str],
    timeout_ms: int = 3000,
) -> Locator | None:
    """Ищет активную кнопку формы, включая Vue-кнопки, отрисованные как span."""
    deadline = asyncio.get_running_loop().time() + max(timeout_ms, 0) / 1000
    while True:
        scopes = [page.locator("form footer"), page.locator("form"), page.locator("body")]
        for scope in scopes:
            controls = scope.locator(
                "button.base-button, a.base-button, span.base-button, "
                "button, a[role='button'], span[role='button']"
            ).filter(has_text=pattern)
            count = await controls.count()
            for index in range(min(count, 20)):
                control = controls.nth(index)
                try:
                    if not await control.is_visible(timeout=0):
                        continue
                    disabled = await control.get_attribute("disabled")
                    aria_disabled = clean_text(
                        await control.get_attribute("aria-disabled") or ""
                    ).casefold()
                    classes = clean_text(await control.get_attribute("class") or "")
                    if (
                        disabled is not None
                        or aria_disabled == "true"
                        or "base-button--disabled" in classes
                    ):
                        continue
                    return control
                except Exception:
                    continue
        if asyncio.get_running_loop().time() >= deadline:
            return None
        await page.wait_for_timeout(100)


async def click_form_action(control: Locator) -> None:
    """Нажимает активную Vue-кнопку независимо от её HTML-тега."""
    await control.scroll_into_view_if_needed()
    try:
        await control.click(timeout=3000)
        return
    except Exception:
        pass
    try:
        await control.dispatch_event("click")
        return
    except Exception:
        pass
    await control.evaluate("element => element.click()")


async def save_draft(page: Page) -> None:
    await dismiss_cookie_banner(page)
    button = await find_enabled_form_action(page, SAVE_BUTTON_PATTERN, 5000)
    if button is None:
        disabled = page.locator("form footer .base-button").filter(
            has_text=SAVE_BUTTON_PATTERN
        )
        if await disabled.count():
            raise AutomationError(
                "Кнопка сохранения найдена, но сейчас отключена сайтом"
            )
        raise AutomationError("Не найдена кнопка сохранения черновика")

    await click_form_action(button)
    spinner = page.locator(
        "[role='progressbar'], [class*='spinner' i], [class*='loader' i]"
    )
    try:
        if await spinner.first.is_visible(timeout=200):
            await spinner.first.wait_for(state="hidden", timeout=10000)
            return
    except Exception:
        pass
    success = page.get_by_text(
        re.compile(r"сохранен|сохранён|изменения сохранены", re.I)
    )
    try:
        await success.first.wait_for(state="visible", timeout=2000)
    except PlaywrightTimeoutError:
        await page.wait_for_timeout(500)


# ---------------------------------------------------------------------------
# Скачивание проекта
# ---------------------------------------------------------------------------


async def download_current_project(
    page: Page,
    candidate: ProjectCandidate,
    settings: Settings,
) -> Path:
    button = await first_visible(
        [
            page.get_by_role("button", name=DOWNLOAD_PROJECT_PATTERN),
            page.get_by_role("link", name=DOWNLOAD_PROJECT_PATTERN),
            page.get_by_text(DOWNLOAD_PROJECT_PATTERN),
        ],
        5000,
    )
    if button is None:
        raise AutomationError(
            "На странице проекта не найдена кнопка «Скачать проект»"
        )

    settings.downloads_dir.mkdir(parents=True, exist_ok=True)
    try:
        async with page.expect_download(timeout=20_000) as download_info:
            await button.scroll_into_view_if_needed()
            await button.click()
        download: Download = await download_info.value
    except PlaywrightTimeoutError as exc:
        raise AutomationError("Сайт не начал скачивание Word-файла") from exc

    suffix = Path(download.suggested_filename or "project.docx").suffix.casefold()
    if suffix not in {".doc", ".docx"}:
        suffix = ".docx"
    destination = unique_path(
        settings.downloads_dir / f"{safe_fs_name(candidate.title)}{suffix}"
    )
    await download.save_as(str(destination))
    return destination


async def process_download_account(
    browser: Browser,
    account: SourceAccount,
    settings: Settings,
    interactive_project_selection: bool,
) -> TransferResult:
    result = TransferResult(
        row_number=account.row_number,
        fio=account.fio,
        operation="скачивание",
        project_name=account.project_hint or "выбор в консоли",
    )
    context: BrowserContext = await browser.new_context(
        viewport={"width": 1360, "height": 900}, accept_downloads=True
    )
    context.set_default_timeout(settings.timeout_ms)
    if settings.block_heavy_resources:
        async def route_handler(route):
            if route.request.resource_type in {"image", "media", "font"}:
                await route.abort()
            else:
                await route.continue_()
        await context.route("**/*", route_handler)
    page = await context.new_page()

    try:
        await login_account(page, account.login, account.password, settings)
        await open_projects(page, settings, require_cards=True)
        await load_all_project_cards(page)
        candidates = await collect_project_candidates(
            page, "" if interactive_project_selection else account.project_hint
        )
        if interactive_project_selection:
            candidate = await choose_candidate_in_console(account, candidates)
        else:
            if not account.project_hint:
                raise AutomationError(
                    "Для автоматического выбора заполните столбец D в исходник.xls "
                    "или запустите программу без параметра --auto-select"
                )
            candidate = await choose_candidate_automatically(
                account, candidates, settings
            )
        result.project_name = candidate.title
        await open_candidate_project(page, candidate)
        downloaded = await download_current_project(page, candidate, settings)
        result.word_file = downloaded.name
        result.status = "успешно"
        return result
    except Exception as exc:
        result.status = "ошибка"
        result.error = str(exc)
        await take_screenshot(
            page, settings, account.row_number, account.fio, "ОШИБКА_СКАЧИВАНИЯ"
        )
        return result
    finally:
        await context.close()


# ---------------------------------------------------------------------------
# Универсальное заполнение полей
# ---------------------------------------------------------------------------


def regex_for_aliases(aliases: Sequence[str], exact: bool = True) -> re.Pattern[str]:
    escaped = [re.escape(alias) for alias in aliases]
    body = "|".join(escaped)
    return re.compile(rf"^\s*(?:{body})\s*$" if exact else rf"(?:{body})", re.I)


async def find_label(
    scope: Page | Locator,
    aliases: Sequence[str],
    timeout_ms: int = 1800,
) -> Locator | None:
    exact_pattern = regex_for_aliases(aliases, exact=True)
    broad_pattern = regex_for_aliases(aliases, exact=False)
    candidates = [
        scope.locator("label").filter(has_text=broad_pattern),
        scope.get_by_text(exact_pattern),
        scope.get_by_text(broad_pattern),
    ]
    return await first_visible(candidates, timeout_ms)


async def control_from_label(label: Locator) -> Locator | None:
    try:
        for_attr = await label.get_attribute("for")
        if for_attr:
            page = label.page
            direct = page.locator(f"#{for_attr}")
            if await direct.count() and await wait_visible(direct, 300):
                return direct.first
    except Exception:
        pass

    following = label.locator(
        "xpath=following::*[self::textarea or self::input or self::select "
        "or @contenteditable='true' or @role='combobox'][1]"
    )
    if await following.count() and await wait_visible(following, 400):
        return following.first

    ancestor = label.locator(
        "xpath=ancestor::*[.//textarea or .//input[not(@type='file')] or .//select "
        "or .//*[@contenteditable='true'] or .//*[@role='combobox']][1]"
    )
    if await ancestor.count():
        controls = ancestor.first.locator(
            "textarea, input:not([type='file']), select, [contenteditable='true'], [role='combobox']"
        )
        count = await controls.count()
        for index in range(min(count, 8)):
            control = controls.nth(index)
            if await wait_visible(control, 150):
                return control
    return None


async def force_fill(control: Locator, value: str) -> None:
    value = clean_text(value)
    if not value:
        return
    tag = (await control.evaluate("el => el.tagName.toLowerCase()")) if await control.count() else ""
    if tag == "select":
        try:
            await control.select_option(label=value)
        except Exception:
            await control.select_option(value=value)
        return
    contenteditable = await control.get_attribute("contenteditable")
    if contenteditable == "true":
        await control.click()
        await control.press("Control+A")
        await control.fill(value)
        return
    try:
        await control.fill(value)
    except Exception:
        await control.evaluate(
            """
            (el, value) => {
                el.removeAttribute('readonly');
                el.removeAttribute('disabled');
                const setter = Object.getOwnPropertyDescriptor(
                    el.tagName === 'TEXTAREA' ? HTMLTextAreaElement.prototype : HTMLInputElement.prototype,
                    'value'
                )?.set;
                if (setter) setter.call(el, value); else el.value = value;
                el.dispatchEvent(new Event('input', {bubbles: true}));
                el.dispatchEvent(new Event('change', {bubbles: true}));
                el.dispatchEvent(new Event('blur', {bubbles: true}));
            }
            """,
            value,
        )


async def fill_by_label(
    scope: Page | Locator,
    aliases: Sequence[str],
    value: str,
    *,
    required: bool = False,
    timeout_ms: int = 1800,
) -> bool:
    if not clean_text(value):
        return False
    label = await find_label(scope, aliases, timeout_ms)
    if label is None:
        if required:
            raise AutomationError(f"Не найдено поле «{aliases[0]}»")
        return False
    control = await control_from_label(label)
    if control is None:
        if required:
            raise AutomationError(f"У поля «{aliases[0]}» не найден элемент ввода")
        return False
    await control.scroll_into_view_if_needed()
    await force_fill(control, value)
    return True


async def select_option_by_label(
    page: Page,
    scope: Page | Locator,
    aliases: Sequence[str],
    value: str,
    *,
    required: bool = False,
) -> bool:
    if not clean_text(value):
        return False
    label = await find_label(scope, aliases, 2000)
    if label is None:
        if required:
            raise AutomationError(f"Не найдено поле выбора «{aliases[0]}»")
        return False
    control = await control_from_label(label)
    if control is None:
        if required:
            raise AutomationError(f"У поля «{aliases[0]}» не найден список")
        return False

    tag = await control.evaluate("el => el.tagName.toLowerCase()")
    if tag == "select":
        try:
            await control.select_option(label=value)
        except Exception:
            await control.select_option(value=value)
        return True

    await control.scroll_into_view_if_needed()
    try:
        await control.click()
    except Exception:
        await control.click(force=True)

    # Если это редактируемый combobox, фильтруем варианты.
    try:
        if tag in {"input", "textarea"}:
            await force_fill(control, value)
    except Exception:
        pass

    option_pattern = re.compile(rf"^\s*{re.escape(value)}\s*$", re.I)
    option = await first_visible(
        [
            page.get_by_role("option", name=option_pattern),
            page.get_by_role("menuitem", name=option_pattern),
            page.get_by_text(option_pattern),
            page.get_by_text(re.compile(re.escape(value), re.I)),
        ],
        2200,
    )
    if option:
        await option.click()
        return True

    # Резерв для автокомплита.
    try:
        await control.press("ArrowDown")
        await control.press("Enter")
        return True
    except Exception:
        if required:
            raise AutomationError(
                f"В поле «{aliases[0]}» не найден вариант «{value}»"
            )
        return False


async def fill_autocomplete_by_label(
    page: Page,
    scope: Page | Locator,
    aliases: Sequence[str],
    value: str,
    *,
    required: bool = False,
) -> bool:
    if not clean_text(value):
        return False
    label = await find_label(scope, aliases, 2000)
    if label is None:
        if required:
            raise AutomationError(f"Не найдено поле «{aliases[0]}»")
        return False
    control = await control_from_label(label)
    if control is None:
        if required:
            raise AutomationError(f"У поля «{aliases[0]}» не найден ввод")
        return False
    await force_fill(control, value)
    await page.wait_for_timeout(250)
    suggestion = await first_visible(
        [
            page.get_by_role("option", name=re.compile(re.escape(value), re.I)),
            page.locator("[role='listbox'] [role='option']"),
            page.locator("[class*='suggest' i], [class*='dropdown' i]").get_by_text(
                re.compile(re.escape(value), re.I)
            ),
        ],
        1200,
    )
    if suggestion:
        await suggestion.first.click()
    else:
        try:
            await control.press("ArrowDown")
            await control.press("Enter")
        except Exception:
            pass
    return True


async def open_tab(page: Page, aliases: Sequence[str]) -> None:
    pattern = regex_for_aliases(aliases, exact=True)
    target = await first_visible(
        [
            page.get_by_role("tab", name=pattern),
            page.get_by_role("button", name=pattern),
            page.get_by_role("link", name=pattern),
            page.get_by_text(pattern),
        ],
        2500,
    )
    if target is None:
        raise AutomationError(f"Не найдена вкладка «{aliases[0]}»")
    await target.scroll_into_view_if_needed()
    await target.click()
    await page.wait_for_timeout(120)


async def fill_date_range(page: Page, start: str, end: str) -> int:
    if not start and not end:
        return 0
    label = await find_label(page, FIELD_ALIASES["dates"], 2200)
    if label is None:
        # Некоторые версии показывают два отдельных поля.
        count = 0
        count += int(
            await fill_by_label(page, FIELD_ALIASES["start_date"], start)
        )
        count += int(await fill_by_label(page, FIELD_ALIASES["end_date"], end))
        return count
    ancestor = label.locator(
        "xpath=ancestor::*[count(.//input[not(@type='file')]) >= 1][1]"
    )
    inputs = ancestor.first.locator("input:not([type='file'])") if await ancestor.count() else page.locator("input")
    count = await inputs.count()
    if count >= 2:
        await force_fill(inputs.nth(0), start)
        await force_fill(inputs.nth(1), end)
        return 2
    control = await control_from_label(label)
    if control:
        await force_fill(control, f"{start} - {end}")
        return 1
    return 0


async def click_add_button(
    scope: Page | Locator,
    aliases: Sequence[str],
    timeout_ms: int = 1800,
) -> bool:
    pattern = regex_for_aliases(aliases, exact=False)
    target = await first_visible(
        [
            scope.get_by_role("button", name=pattern),
            scope.get_by_role("link", name=pattern),
            scope.get_by_text(pattern),
        ],
        timeout_ms,
    )
    if target is None:
        return False
    await target.scroll_into_view_if_needed()
    await target.click()
    return True


async def nearest_container_with_texts(
    locator: Locator, texts: Sequence[str]
) -> Locator | None:
    conditions = " and ".join(
        f".//*[contains(normalize-space(), {xpath_literal(text)})]" for text in texts
    )
    ancestor = locator.locator(f"xpath=ancestor::*[{conditions}][1]")
    if await ancestor.count():
        return ancestor.first
    return None


def xpath_literal(value: str) -> str:
    if "'" not in value:
        return f"'{value}'"
    if '"' not in value:
        return f'"{value}"'
    parts = value.split("'")
    return "concat(" + ", \"'\", ".join(f"'{part}'" for part in parts) + ")"


# ---------------------------------------------------------------------------
# Создание нового черновика
# ---------------------------------------------------------------------------


async def open_new_project_template(page: Page, settings: Settings) -> None:
    """Открывает реальную стартовую форму шаблона без создания черновика."""
    await open_projects(page, settings, require_cards=False)
    await dismiss_cookie_banner(page)

    add_button = await first_visible(
        [
            page.get_by_role("button", name=ADD_PROJECT_PATTERN),
            page.get_by_role("link", name=ADD_PROJECT_PATTERN),
            page.get_by_text(ADD_PROJECT_PATTERN),
        ],
        5000,
    )
    if add_button is None:
        raise AutomationError("Не найдена кнопка «Добавить проект»")
    try:
        await add_button.click(timeout=5000)
    except Exception:
        await add_button.evaluate("element => element.click()")

    modal = page.locator(".modal-vue--content").filter(
        has_text=re.compile(r"Выберите шаблон проекта", re.I)
    )
    try:
        await modal.first.wait_for(state="visible", timeout=7000)
    except PlaywrightTimeoutError as exc:
        raise AutomationError(
            "После нажатия «Добавить проект» не открылось окно выбора шаблона"
        ) from exc
    modal = modal.first

    template_pattern = re.compile(
        rf"^\s*{re.escape(settings.project_template_name)}\s*$", re.I
    )
    template = modal.locator(
        "a[href^='/projects/create/'], a[href*='/projects/create/']"
    ).filter(has_text=template_pattern)

    if not await wait_visible(template, 5000):
        search = modal.locator("input#search, input[placeholder='Поиск']")
        if await wait_visible(search, 1000):
            await search.first.fill(settings.project_template_name)
            await page.wait_for_timeout(350)
        if not await wait_visible(template, 3000):
            raise AutomationError(
                "Не найден шаблон проекта в окне выбора: "
                + settings.project_template_name
            )

    href = clean_text(await template.first.get_attribute("href") or "")
    if not href:
        raise AutomationError("У ссылки шаблона отсутствует адрес создания проекта")

    await page.goto(
        urljoin(page.url, href),
        wait_until="domcontentloaded",
        timeout=max(settings.timeout_ms, 30_000),
    )
    await dismiss_cookie_banner(page)

    try:
        await page.locator("form input#projectName").first.wait_for(
            state="visible", timeout=10000
        )
    except PlaywrightTimeoutError as exc:
        raise AutomationError(
            "После выбора шаблона не открылась стартовая форма с полем «Название проекта»"
        ) from exc


async def fill_initial_project_title(page: Page, title: str) -> None:
    """Заполняет именно стартовое поле #projectName и будит Vue-валидацию."""
    title = clean_text(title)
    if not title:
        raise AutomationError("Название проекта отсутствует в Word-файле")

    control = page.locator("form input#projectName").first
    try:
        await control.wait_for(state="visible", timeout=5000)
    except PlaywrightTimeoutError as exc:
        raise AutomationError("Не найдено стартовое поле #projectName") from exc

    await control.click()
    await control.fill("")
    await control.fill(title)
    await control.dispatch_event("input")
    await control.dispatch_event("change")
    try:
        await control.press("Tab")
    except Exception:
        await control.dispatch_event("blur")
    await page.wait_for_timeout(200)

    actual = clean_text(await control.input_value())
    if actual != title:
        raise AutomationError(
            f"Название проекта не записалось в поле: ожидалось «{title}», получено «{actual}»"
        )


async def select_initial_project_region(page: Page, region: str) -> None:
    """Работает с реальным Vue multiselect стартового поля региона."""
    region = clean_text(region)
    if not region:
        raise AutomationError(
            "В Word-файле отсутствует регион реализации проекта"
        )

    field = page.locator(
        "form div.base-select[title='Регион реализации проекта']"
    ).first
    try:
        await field.wait_for(state="visible", timeout=5000)
    except PlaywrightTimeoutError as exc:
        raise AutomationError(
            "Не найден стартовый список «Регион реализации проекта»"
        ) from exc

    selected = field.locator(".multiselect__single")
    if await selected.count():
        current = clean_text(await selected.first.inner_text())
        if normalize_text(current) == normalize_text(region):
            return

    tags = field.locator(".multiselect__tags").first
    input_box = field.locator("input.multiselect__input").first
    try:
        await tags.click(force=True)
    except Exception:
        await field.locator(".multiselect").first.click(force=True)

    try:
        await input_box.fill(region)
    except Exception:
        await input_box.evaluate(
            """
            (el, value) => {
                el.removeAttribute('disabled');
                el.style.width = '100%';
                el.value = value;
                el.dispatchEvent(new Event('input', {bubbles: true}));
                el.dispatchEvent(new Event('change', {bubbles: true}));
            }
            """,
            region,
        )
    await page.wait_for_timeout(400)

    exact = re.compile(rf"^\s*{re.escape(region)}\s*$", re.I)
    option = page.locator(".multiselect__option").filter(has_text=exact)
    chosen = await first_visible([option], 5000)
    if chosen is not None:
        try:
            await chosen.click(timeout=3000)
        except Exception:
            await chosen.evaluate("element => element.click()")
    else:
        try:
            await input_box.press("ArrowDown")
            await input_box.press("Enter")
        except Exception as exc:
            raise AutomationError(
                f"В списке регионов не найден вариант «{region}»"
            ) from exc

    await page.wait_for_timeout(250)
    if await selected.count():
        current = clean_text(await selected.first.inner_text())
        if normalize_text(current) == normalize_text(region):
            return
    raise AutomationError(f"Регион «{region}» не был выбран в форме")


async def click_create_draft(page: Page, settings: Settings) -> None:
    """Нажимает активную кнопку «Создать черновик» и ждёт режим редактирования."""
    await dismiss_cookie_banner(page)
    create = await find_enabled_form_action(
        page, re.compile(r"^\s*Создать черновик\s*$", re.I), 12000
    )
    if create is None:
        disabled = page.locator("form footer .base-button").filter(
            has_text=re.compile(r"^\s*Создать черновик\s*$", re.I)
        )
        if await disabled.count():
            raise AutomationError(
                "Кнопка «Создать черновик» найдена, но остаётся отключённой. "
                "Проверьте название проекта и выбранный регион"
            )
        raise AutomationError("Не найдена кнопка «Создать черновик»")

    old_url = page.url
    await click_form_action(create)

    deadline = asyncio.get_running_loop().time() + max(
        settings.timeout_ms, 30_000
    ) / 1000
    while asyncio.get_running_loop().time() < deadline:
        await dismiss_cookie_banner(page)
        save = await find_enabled_form_action(page, SAVE_BUTTON_PATTERN, 0)
        if save is not None:
            return

        error_message = page.locator(
            ".vue-notification.error, .vue-notification-template.error, "
            ".vue-notification-wrapper"
        )
        if await error_message.count():
            try:
                text = clean_text(await error_message.last.inner_text())
                if text:
                    raise AutomationError(
                        f"Сайт отклонил создание черновика: {text}"
                    )
            except AutomationError:
                raise
            except Exception:
                pass

        await page.wait_for_timeout(250)

    location_note = f" Текущий адрес: {page.url}" if page.url != old_url else ""
    raise AutomationError(
        "После нажатия «Создать черновик» не появился режим сохранения."
        + location_note
    )


async def run_optional_website_step(
    step_name: str,
    operation,
    skipped: list[str],
) -> tuple[bool, object | None]:
    """Повторяет либо пропускает только текущий шаг, не весь аккаунт."""
    while True:
        try:
            return True, await operation()
        except KeyboardInterrupt:
            raise
        except Exception as exc:
            action = await ask_error_action(
                step_name,
                exc,
                skip_label="пропустить только этот шаг",
            )
            if action == "retry":
                continue
            if action == "skip":
                skipped.append(f"Шаг «{step_name}» пропущен: {exc}")
                return False, None
            raise KeyboardInterrupt


async def create_initial_draft(
    page: Page,
    project: ProjectData,
    settings: Settings,
    skipped: list[str],
) -> tuple[int, bool, bool]:
    """Выполняет стартовые действия отдельно, с пропуском каждого шага."""
    filled = 0

    title_ok, _ = await run_optional_website_step(
        "заполнение названия проекта",
        lambda: fill_initial_project_title(page, project.title),
        skipped,
    )
    if title_ok:
        filled += 1

    region_ok, _ = await run_optional_website_step(
        "выбор региона реализации проекта",
        lambda: select_initial_project_region(page, project.region),
        skipped,
    )
    if region_ok:
        filled += 1

    draft_ok, _ = await run_optional_website_step(
        "нажатие кнопки «Создать черновик»",
        lambda: click_create_draft(page, settings),
        skipped,
    )
    return filled, title_ok and region_ok, draft_ok


async def fill_general_tab(
    page: Page,
    project: ProjectData,
    settings: Settings,
    *,
    identity_already_filled: bool = False,
) -> tuple[int, list[str]]:
    await open_tab(page, ("Общее",))
    filled = 0
    skipped: list[str] = []
    required = settings.strict_required_fields

    if not identity_already_filled:
        filled += int(
            await fill_by_label(
                page, FIELD_ALIASES["title"], project.title, required=True
            )
        )
        if project.region:
            ok = await select_option_by_label(
                page,
                page,
                FIELD_ALIASES["region"],
                project.region,
                required=required,
            )
            if not ok:
                ok = await fill_autocomplete_by_label(
                    page,
                    page,
                    FIELD_ALIASES["region"],
                    project.region,
                    required=required,
                )
            filled += int(ok)
        else:
            skipped.append("Регион реализации проекта отсутствует в Word")

    if project.scale:
        filled += int(
            await select_option_by_label(
                page, page, FIELD_ALIASES["scale"], project.scale, required=required
            )
        )
    else:
        skipped.append("Масштаб проекта отсутствует в Word")

    date_count = await fill_date_range(page, project.start_date, project.end_date)
    filled += date_count
    if not date_count:
        skipped.append("Сроки проекта не заполнены")

    skipped.extend(
        [
            "Логотип не переносится из Word",
            "Личные данные автора и резюме пропущены по настройке",
        ]
    )
    return filled, skipped


async def fill_about_tab(
    page: Page, project: ProjectData, settings: Settings
) -> tuple[int, list[str]]:
    await open_tab(page, ("О проекте", "О проекте "))
    required = settings.strict_required_fields
    filled = 0
    skipped: list[str] = []
    fields = (
        ("summary", project.summary),
        ("target_groups", project.target_groups),
        ("problem", project.problem),
        ("goal", project.goal),
        ("team_experience", project.team_experience),
        ("development", project.development),
    )
    for key, value in fields:
        if value:
            filled += int(
                await fill_by_label(
                    page, FIELD_ALIASES[key], value, required=required
                )
            )
        else:
            skipped.append(f"В Word отсутствует: {FIELD_ALIASES[key][0]}")

    if project.geography:
        first = project.geography[0]
        ok = await fill_autocomplete_by_label(
            page, page, FIELD_ALIASES["geography"], first, required=required
        )
        filled += int(ok)
        for location in project.geography[1:]:
            if not await click_add_button(page, ("Добавить поле", "Добавить географию")):
                skipped.append(f"Не удалось добавить географию: {location}")
                continue
            labels = page.get_by_text(regex_for_aliases(FIELD_ALIASES["geography"], True))
            count = await labels.count()
            scope: Page | Locator = page
            if count:
                label = labels.nth(count - 1)
                container = label.locator(
                    "xpath=ancestor::*[.//input[not(@type='file')]][1]"
                )
                if await container.count():
                    scope = container.first
            ok = await fill_autocomplete_by_label(
                page, scope, FIELD_ALIASES["geography"], location
            )
            filled += int(ok)
    else:
        skipped.append("География проекта отсутствует в Word")

    skipped.extend(
        [
            "Файл, подтверждающий проблему, не переносится из Word",
            "Регион нахождения автора пропущен как персональное поле",
        ]
    )
    return filled, skipped


async def fill_results_tab(
    page: Page, project: ProjectData, settings: Settings
) -> tuple[int, list[str]]:
    await open_tab(page, ("Результаты",))
    required = settings.strict_required_fields
    filled = 0
    skipped: list[str] = []
    values = (
        ("events_count", project.events_count),
        ("publications_count", project.publications_count),
        ("views_count", project.views_count),
        ("social_effect", project.social_effect),
    )
    for key, value in values:
        if value:
            fill_value = format_number(value) if key != "social_effect" else value
            filled += int(
                await fill_by_label(
                    page, FIELD_ALIASES[key], fill_value, required=required
                )
            )
        else:
            skipped.append(f"В Word отсутствует: {FIELD_ALIASES[key][0]}")
    return filled, skipped


async def task_blocks(page: Page) -> list[Locator]:
    labels = page.get_by_text(regex_for_aliases(CALENDAR_ALIASES["task"], True))
    count = await labels.count()
    result: list[Locator] = []
    seen: set[str] = set()
    for index in range(count):
        label = labels.nth(index)
        block = label.locator(
            "xpath=ancestor::*[.//*[contains(normalize-space(),'Название мероприятия')] "
            "and .//*[contains(normalize-space(),'Добавить мероприятие')]][1]"
        )
        if await block.count() == 0:
            block = label.locator(
                "xpath=ancestor::*[.//*[contains(normalize-space(),'Название мероприятия')]][1]"
            )
        if await block.count():
            marker = await block.first.evaluate(
                "el => el.dataset.transferId || (el.dataset.transferId = Math.random().toString(36))"
            )
            if marker not in seen:
                seen.add(marker)
                result.append(block.first)
    return result


async def event_blocks(task_block: Locator) -> list[Locator]:
    labels = task_block.get_by_text(regex_for_aliases(CALENDAR_ALIASES["title"], True))
    count = await labels.count()
    result: list[Locator] = []
    seen: set[str] = set()
    for index in range(count):
        label = labels.nth(index)
        block = label.locator(
            "xpath=ancestor::*[.//*[contains(normalize-space(),'Описание мероприятия')] "
            "and .//*[contains(normalize-space(),'Крайняя дата')]][1]"
        )
        if await block.count():
            marker = await block.first.evaluate(
                "el => el.dataset.transferEventId || (el.dataset.transferEventId = Math.random().toString(36))"
            )
            if marker not in seen:
                seen.add(marker)
                result.append(block.first)
    return result


async def get_or_create_task_block(page: Page, index: int) -> Locator:
    blocks = await task_blocks(page)
    while len(blocks) <= index:
        if not await click_add_button(page, ("Добавить задачу",)):
            raise AutomationError("Не найдена кнопка «Добавить задачу»")
        await page.wait_for_timeout(150)
        blocks = await task_blocks(page)
    return blocks[index]


async def get_or_create_event_block(task_block: Locator, index: int) -> Locator:
    blocks = await event_blocks(task_block)
    while len(blocks) <= index:
        if not await click_add_button(
            task_block, ("Добавить мероприятие в задачу", "Добавить мероприятие")
        ):
            raise AutomationError("Не найдена кнопка добавления мероприятия")
        await task_block.page.wait_for_timeout(150)
        blocks = await event_blocks(task_block)
    return blocks[index]


async def fill_calendar_tab(
    page: Page, project: ProjectData, settings: Settings
) -> tuple[int, list[str]]:
    await open_tab(page, ("Календарный план",))
    if not project.calendar:
        return 0, ["Календарный план не распознан в Word"]

    grouped: list[tuple[str, list[CalendarEvent]]] = []
    for event in project.calendar:
        task = event.task or "Реализовать комплекс мероприятий проекта"
        if grouped and normalize_text(grouped[-1][0]) == normalize_text(task):
            grouped[-1][1].append(event)
        else:
            grouped.append((task, [event]))

    filled = 0
    skipped: list[str] = []
    for task_index, (task_text, events) in enumerate(grouped):
        task_block = await get_or_create_task_block(page, task_index)
        if await fill_by_label(
            task_block, CALENDAR_ALIASES["task"], task_text, required=True
        ):
            filled += 1
        for event_index, event in enumerate(events):
            event_block = await get_or_create_event_block(task_block, event_index)
            mapping = (
                ("title", event.title, True),
                ("deadline", event.deadline, True),
                ("description", event.description, True),
                ("unique_participants", format_number(event.unique_participants), False),
                ("repeat_participants", format_number(event.repeat_participants), False),
                ("publications", format_number(event.publications), False),
                ("views", format_number(event.views), False),
                ("additional", event.additional, False),
            )
            for key, value, mandatory in mapping:
                if not value and mandatory:
                    skipped.append(
                        f"Мероприятие «{event.title or event_index + 1}»: отсутствует {CALENDAR_ALIASES[key][0]}"
                    )
                    continue
                if value:
                    ok = await fill_by_label(
                        event_block,
                        CALENDAR_ALIASES[key],
                        value,
                        required=mandatory and settings.strict_required_fields,
                    )
                    filled += int(ok)
    return filled, skipped


async def cofinancing_partner_cards(page: Page) -> list[Locator]:
    labels = page.get_by_text(regex_for_aliases(("Название партнера", "Название партнёра"), True))
    count = await labels.count()
    result: list[Locator] = []
    seen: set[str] = set()
    for index in range(count):
        label = labels.nth(index)
        block = label.locator(
            "xpath=ancestor::*[.//*[contains(normalize-space(),'Перечень расходов')] "
            "and .//*[contains(normalize-space(),'Сумма')]][1]"
        )
        if await block.count():
            marker = await block.first.evaluate(
                "el => el.dataset.transferPartnerId || (el.dataset.transferPartnerId = Math.random().toString(36))"
            )
            if marker not in seen:
                seen.add(marker)
                result.append(block.first)
    return result


async def fill_cofinancing_tab(
    page: Page, project: ProjectData, settings: Settings
) -> tuple[int, list[str]]:
    await open_tab(page, ("Софинансирование",))
    filled = 0
    skipped = ["Подтверждающие файлы софинансирования не переносятся из Word"]
    own = [entry for entry in project.cofinancing if entry.kind == "own"]
    partners = [entry for entry in project.cofinancing if entry.kind == "partner"]

    if own:
        combined_expenses = "\n".join(entry.expenses for entry in own if entry.expenses)
        amounts = [float(format_number(entry.amount)) for entry in own if entry.amount]
        combined_amount = str(sum(amounts)).rstrip("0").rstrip(".") if amounts else ""
        own_heading = await find_label(page, ("Собственные средства",), 1800)
        own_scope: Page | Locator = page
        if own_heading:
            block = own_heading.locator(
                "xpath=ancestor::*[.//*[contains(normalize-space(),'Перечень расходов')] "
                "and .//*[contains(normalize-space(),'Сумма')]][1]"
            )
            if await block.count():
                own_scope = block.first
        filled += int(
            await fill_by_label(
                own_scope, COFINANCING_ALIASES["expenses"], combined_expenses
            )
        )
        filled += int(
            await fill_by_label(
                own_scope, COFINANCING_ALIASES["amount"], combined_amount
            )
        )

    for index, entry in enumerate(partners):
        cards = await cofinancing_partner_cards(page)
        while len(cards) <= index:
            if not await click_add_button(page, ("Добавить поле",)):
                skipped.append(
                    f"Не удалось добавить партнёра «{entry.partner_name or index + 1}»"
                )
                break
            await page.wait_for_timeout(150)
            cards = await cofinancing_partner_cards(page)
        if len(cards) <= index:
            continue
        card = cards[index]
        if entry.partner_name:
            filled += int(
                await fill_by_label(
                    card, COFINANCING_ALIASES["partner_name"], entry.partner_name
                )
            )
        if entry.support_type:
            filled += int(
                await select_option_by_label(
                    page,
                    card,
                    COFINANCING_ALIASES["support_type"],
                    entry.support_type,
                )
            )
        if entry.expenses:
            filled += int(
                await fill_by_label(
                    card, COFINANCING_ALIASES["expenses"], entry.expenses
                )
            )
        if entry.amount:
            filled += int(
                await fill_by_label(
                    card,
                    COFINANCING_ALIASES["amount"],
                    format_number(entry.amount),
                )
            )
    return filled, skipped


async def expense_category_container(page: Page, category: str) -> Locator | None:
    pattern = re.compile(rf"^\s*{re.escape(category)}\s*$", re.I)
    heading = await first_visible(
        [
            page.get_by_role("button", name=pattern),
            page.get_by_text(pattern),
            page.get_by_text(re.compile(re.escape(category), re.I)),
        ],
        1800,
    )
    if heading is None:
        return None
    block = heading.locator(
        "xpath=ancestor::*[.//*[contains(normalize-space(),'Добавить строку услуги')] "
        "or .//*[contains(normalize-space(),'Услуга / Товар')]][1]"
    )
    if await block.count():
        return block.first
    try:
        await heading.click()
        await page.wait_for_timeout(120)
    except Exception:
        pass
    block = heading.locator(
        "xpath=ancestor::*[.//*[contains(normalize-space(),'Добавить строку услуги')] "
        "or .//*[contains(normalize-space(),'Услуга / Товар')]][1]"
    )
    return block.first if await block.count() else None


async def expense_rows(container: Locator) -> list[Locator]:
    labels = container.get_by_text(regex_for_aliases(EXPENSE_ALIASES["item"], True))
    count = await labels.count()
    rows: list[Locator] = []
    seen: set[str] = set()
    for index in range(count):
        label = labels.nth(index)
        block = label.locator(
            "xpath=ancestor::*[.//*[contains(normalize-space(),'Цена')] "
            "and .//*[contains(normalize-space(),'Кол-во')]][1]"
        )
        if await block.count():
            marker = await block.first.evaluate(
                "el => el.dataset.transferExpenseId || (el.dataset.transferExpenseId = Math.random().toString(36))"
            )
            if marker not in seen:
                seen.add(marker)
                rows.append(block.first)
    return rows


async def fill_expenses_tab(
    page: Page, project: ProjectData, settings: Settings
) -> tuple[int, list[str]]:
    await open_tab(page, ("Расходы",))
    if not project.expenses:
        return 0, ["Строки сметы не распознаны в Word"]
    filled = 0
    skipped: list[str] = []
    by_category: dict[str, list[ExpenseItem]] = {}
    for item in project.expenses:
        if not item.category:
            skipped.append(f"Не определена категория расхода: {item.item}")
            continue
        by_category.setdefault(item.category, []).append(item)

    for category, items in by_category.items():
        container = await expense_category_container(page, category)
        if container is None:
            skipped.append(f"На сайте не найдена категория расхода: {category}")
            continue
        rows = await expense_rows(container)
        for index, item in enumerate(items):
            while len(rows) <= index:
                if not await click_add_button(
                    container, ("Добавить строку услуги / товара", "Добавить строку")
                ):
                    skipped.append(f"Не удалось добавить строку сметы: {item.item}")
                    break
                await page.wait_for_timeout(120)
                rows = await expense_rows(container)
            if len(rows) <= index:
                continue
            row = rows[index]
            mapping = (
                ("item", item.item),
                ("justification", item.justification),
                ("price", format_number(item.price)),
                ("quantity", format_number(item.quantity, "1")),
            )
            for key, value in mapping:
                if value:
                    ok = await fill_by_label(
                        row,
                        EXPENSE_ALIASES[key],
                        value,
                        required=(
                            key in {"item", "price", "quantity"}
                            and settings.strict_required_fields
                        ),
                    )
                    filled += int(ok)
    return filled, skipped


async def fill_project(
    page: Page,
    project: ProjectData,
    settings: Settings,
    *,
    identity_already_filled: bool = False,
) -> tuple[int, list[str]]:
    total = 0
    skipped: list[str] = []

    async def general_stage(
        current_page: Page,
        current_project: ProjectData,
        current_settings: Settings,
    ) -> tuple[int, list[str]]:
        return await fill_general_tab(
            current_page,
            current_project,
            current_settings,
            identity_already_filled=identity_already_filled,
        )

    stages = (
        ("Общее", general_stage),
        ("О проекте", fill_about_tab),
        ("Результаты", fill_results_tab),
        ("Календарный план", fill_calendar_tab),
        ("Софинансирование", fill_cofinancing_tab),
        ("Расходы", fill_expenses_tab),
    )

    changed = False
    for tab_name, function in stages:
        while True:
            try:
                count, tab_skipped = await function(page, project, settings)
                total += count
                skipped.extend(tab_skipped)
                changed = changed or count > 0

                if count > 0 and settings.save_after_each_tab:
                    saved, _ = await run_optional_website_step(
                        f"сохранение вкладки «{tab_name}»",
                        lambda: save_draft(page),
                        skipped,
                    )
                    if saved:
                        changed = False
                break
            except KeyboardInterrupt:
                raise
            except Exception as exc:
                action = await ask_error_action(
                    f"заполнение вкладки «{tab_name}»",
                    exc,
                    skip_label="пропустить только эту вкладку",
                )
                if action == "retry":
                    continue
                if action == "skip":
                    skipped.append(f"Вкладка «{tab_name}» пропущена: {exc}")
                    break
                raise KeyboardInterrupt

    skipped.extend(
        [
            "Вкладка «Команда» пропущена по заданию",
            "Вкладка «Медиа» пропущена по заданию",
            "Вкладка «Доп. файлы»: бинарные вложения отсутствуют внутри Word-выгрузки",
        ]
    )
    if project.extra_file_descriptions:
        skipped.append(
            "Описания дополнительных файлов найдены, но сами файлы невозможно восстановить из Word: "
            + ", ".join(project.extra_file_descriptions)
        )

    if changed or not settings.save_after_each_tab:
        await run_optional_website_step(
            "итоговое сохранение черновика",
            lambda: save_draft(page),
            skipped,
        )
    return total, skipped



# ---------------------------------------------------------------------------
# API-заполнение проекта
# ---------------------------------------------------------------------------


class MyRosmolApiError(AutomationError):
    """Ошибка API ФГАИС с расшифровкой полей формы."""


def _schema_label_map(template: dict[str, Any]) -> dict[str, str]:
    labels: dict[str, str] = {}

    def visit(field: dict[str, Any]) -> None:
        field_id = clean_text(str(field.get("ID", "")))
        if field_id:
            labels[field_id] = clean_text(str(field.get("label", ""))) or field_id
        for part in field.get("parts") or []:
            if isinstance(part, dict):
                visit(part)

    form = template.get("form") or {}
    for tab in form.get("tabs") or []:
        for group in tab.get("groups") or []:
            for field_info in group.get("fields") or []:
                if isinstance(field_info, dict):
                    visit(field_info)
    return labels


def _format_api_errors(payload: Any, labels: dict[str, str] | None = None) -> str:
    labels = labels or {}
    if isinstance(payload, str):
        return clean_text(payload)

    parts: list[str] = []
    if isinstance(payload, dict):
        message = clean_text(str(payload.get("message", "")))
        if message:
            parts.append(message)
        errors = payload.get("errors")
        if errors is not None:
            def walk(value: Any, prefix: str = "") -> None:
                if isinstance(value, dict):
                    for key, child in value.items():
                        visible = labels.get(str(key), str(key))
                        next_prefix = f"{prefix} / {visible}" if prefix else visible
                        walk(child, next_prefix)
                elif isinstance(value, list):
                    for child in value:
                        walk(child, prefix)
                elif value not in (None, ""):
                    text = clean_text(str(value))
                    if prefix:
                        parts.append(f"{prefix}: {text}")
                    else:
                        parts.append(text)
            walk(errors)
    elif payload not in (None, ""):
        parts.append(clean_text(str(payload)))

    unique: list[str] = []
    seen: set[str] = set()
    for item in parts:
        if item and item not in seen:
            seen.add(item)
            unique.append(item)
    return "; ".join(unique) or "Неизвестная ошибка API"


def _extract_token_candidate(value: Any) -> str:
    if isinstance(value, str):
        text = value.strip()
        if text.casefold().startswith("bearer "):
            return text
        if text.count(".") == 2 and len(text) > 80:
            return f"Bearer {text}"
        try:
            parsed = json.loads(text)
        except Exception:
            return ""
        return _extract_token_candidate(parsed)
    if isinstance(value, dict):
        priority = (
            "authorization",
            "accessToken",
            "access_token",
            "token",
            "jwt",
        )
        for key in priority:
            if key in value:
                token = _extract_token_candidate(value[key])
                if token:
                    return token
        for child in value.values():
            token = _extract_token_candidate(child)
            if token:
                return token
    if isinstance(value, list):
        for child in value:
            token = _extract_token_candidate(child)
            if token:
                return token
    return ""


async def resolve_authorization_header(
    page: Page,
    captured: dict[str, str],
    settings: Settings,
) -> str:
    for _ in range(30):
        value = clean_text(captured.get("authorization", ""))
        if value:
            return value
        await page.wait_for_timeout(100)

    try:
        storage = await page.evaluate(
            """
            () => ({
                local: Object.fromEntries(
                    Array.from({length: localStorage.length}, (_, index) => {
                        const key = localStorage.key(index);
                        return [key, localStorage.getItem(key)];
                    })
                ),
                session: Object.fromEntries(
                    Array.from({length: sessionStorage.length}, (_, index) => {
                        const key = sessionStorage.key(index);
                        return [key, sessionStorage.getItem(key)];
                    })
                )
            })
            """
        )
        token = _extract_token_candidate(storage)
        if token:
            return token
    except Exception:
        pass

    # Дополнительный переход инициирует профильные API-запросы приложения,
    # из которых обработчик request получает реальный Authorization.
    try:
        await page.goto(
            settings.projects_url,
            wait_until="domcontentloaded",
            timeout=max(settings.timeout_ms, 30_000),
        )
        await page.wait_for_timeout(1200)
    except Exception:
        pass
    value = clean_text(captured.get("authorization", ""))
    if value:
        return value

    raise AutomationError(
        "Не удалось получить токен авторизации после входа. "
        "Откройте страницу «Мои проекты» в появившемся браузере и повторите шаг."
    )


async def api_json_request(
    context: BrowserContext,
    settings: Settings,
    authorization: str,
    method: str,
    path: str,
    *,
    payload: dict[str, Any] | None = None,
    labels: dict[str, str] | None = None,
) -> Any:
    url = path if path.startswith(("http://", "https://")) else (
        f"{settings.base_url}/api/{path.lstrip('/')}"
    )
    headers = {
        "Accept": "application/json, text/plain, */*",
        "Authorization": authorization,
    }
    data: str | None = None
    if payload is not None:
        headers["Content-Type"] = "application/json;charset=UTF-8"
        data = json.dumps(payload, ensure_ascii=False, separators=(",", ":"))

    response = await context.request.fetch(
        url,
        method=method.upper(),
        headers=headers,
        data=data,
        fail_on_status_code=False,
        timeout=max(settings.timeout_ms, 60_000),
    )
    text = await response.text()
    try:
        body: Any = json.loads(text) if text else None
    except Exception:
        body = text

    if not response.ok:
        raise MyRosmolApiError(
            f"API {method.upper()} {url}: HTTP {response.status}. "
            f"{_format_api_errors(body, labels)}"
        )
    return body



def _existing_project_status(project: dict[str, Any]) -> str:
    if bool(project.get("isArchive")):
        return "архив"
    if bool(project.get("isDraft")):
        return "черновик"
    return "подан / не черновик"


def _existing_project_is_editable(project: dict[str, Any]) -> bool:
    return (
        bool(project.get("isDraft"))
        and bool(project.get("isManageable"))
        and not bool(project.get("isArchive"))
        and bool(clean_text(str(project.get("ID", ""))))
    )


def _existing_project_line(index: int, project: dict[str, Any]) -> str:
    details = [_existing_project_status(project)]
    created_at = clean_text(str(project.get("createdAt", "")))
    if created_at:
        details.append(created_at)
    template = project.get("template") or {}
    template_name = clean_text(
        str(template.get("displayName") or template.get("templateName") or "")
    )
    if template_name:
        details.append(template_name)
    if not bool(project.get("isManageable")):
        details.append("нет права редактирования")
    return (
        f"{index}. {clean_text(str(project.get('name', ''))) or 'Без названия'} "
        f"({'; '.join(details)})"
    )


async def get_account_projects_api(
    context: BrowserContext,
    settings: Settings,
    authorization: str,
) -> list[dict[str, Any]]:
    """Получает все неархивные проекты текущего аккаунта через API."""
    result: list[dict[str, Any]] = []
    page_number = 1
    page_size = 100
    fields = (
        "isManageable,feedbackComment,ID,name,isArchive,isDraft,"
        "feedbackRequestAvailable,feedbackRequested,template,createdAt,region,"
        "consultation,template.displayName,template.templateName,template.type.code"
    )
    while page_number <= 100:
        items = await api_json_request(
            context,
            settings,
            authorization,
            "GET",
            "/projects"
            f"?my=1&isProgram=0&fields={fields}&isArchive=false"
            f"&page={page_number}&per-page={page_size}",
        )
        if not isinstance(items, list):
            raise AutomationError("API вернул некорректный список проектов аккаунта")
        result.extend(item for item in items if isinstance(item, dict))
        if len(items) < page_size:
            break
        page_number += 1
    return result


def find_existing_projects_by_title(
    projects: Sequence[dict[str, Any]],
    title: str,
) -> list[dict[str, Any]]:
    expected = normalize_text(title)
    return [
        project
        for project in projects
        if not bool(project.get("isArchive"))
        and normalize_text(str(project.get("name", ""))) == expected
    ]


async def choose_existing_project_or_new(
    projects: Sequence[dict[str, Any]],
    title: str,
) -> tuple[str, dict[str, Any] | None]:
    """Возвращает ('existing', проект) либо ('new', None)."""
    matches = find_existing_projects_by_title(projects, title)
    if not matches:
        return "new", None

    print(f"\nВ аккаунте уже найден проект с названием «{title}».")
    for index, item in enumerate(matches, 1):
        print(_existing_project_line(index, item))

    editable = [item for item in matches if _existing_project_is_editable(item)]
    if not editable:
        print(
            "Совпадающие проекты нельзя редактировать: среди них нет доступного "
            "черновика. Будет создан новый проект."
        )
        return "new", None

    while True:
        try:
            raw = (
                await asyncio.to_thread(
                    input,
                    "1 - заполнить существующий черновик; "
                    "2 - создать новый проект [Enter = 1]: ",
                )
            ).strip().casefold()
        except EOFError as exc:
            raise AutomationError(
                "Не удалось запросить способ обработки совпадающего проекта"
            ) from exc

        if raw in {"2", "n", "т", "new", "новый", "создать"}:
            return "new", None
        if raw not in {"", "1", "e", "у", "existing", "существующий", "заполнить"}:
            print("Введите 1 или 2.")
            continue

        if len(editable) == 1:
            selected = editable[0]
            print(
                "Будет заполнен существующий черновик: "
                f"{clean_text(str(selected.get('name', '')))} "
                f"(ID {selected.get('ID')})."
            )
            return "existing", selected

        print("Выберите редактируемый черновик:")
        editable_by_number: dict[int, dict[str, Any]] = {}
        for display_index, item in enumerate(matches, 1):
            if _existing_project_is_editable(item):
                editable_by_number[display_index] = item
                print(_existing_project_line(display_index, item))
        while True:
            try:
                selected_raw = (
                    await asyncio.to_thread(
                        input,
                        "Введите номер существующего черновика "
                        "(B - вернуться к выбору действия): ",
                    )
                ).strip().casefold()
            except EOFError as exc:
                raise AutomationError(
                    "Не удалось прочитать номер существующего черновика"
                ) from exc
            if selected_raw in {"b", "и", "back", "назад"}:
                break
            if selected_raw.isdigit() and int(selected_raw) in editable_by_number:
                return "existing", editable_by_number[int(selected_raw)]
            print("Введите номер доступного черновика из списка или B.")


async def get_project_detail_api(
    context: BrowserContext,
    settings: Settings,
    authorization: str,
    project_id: str,
) -> dict[str, Any]:
    detailed = await api_json_request(
        context,
        settings,
        authorization,
        "GET",
        f"/projects/{project_id}"
        "?expand=agreement,agreementStatusCode,form,costs,won,wonAmount,"
        "programInfo,isAllFieldsDisabled",
    )
    if not isinstance(detailed, dict) or not detailed.get("ID"):
        raise AutomationError("API не вернул данные существующего проекта")
    return detailed


async def get_project_template_by_id_api(
    context: BrowserContext,
    settings: Settings,
    authorization: str,
    template_id: str,
) -> dict[str, Any]:
    template_id = clean_text(template_id)
    if not template_id:
        raise AutomationError("У проекта отсутствует ID шаблона")
    detailed = await api_json_request(
        context,
        settings,
        authorization,
        "GET",
        f"/project-templates/{template_id}"
        "?expand=form,reportForm,costCategories,costSettings,projectChildTemplates",
    )
    if not isinstance(detailed, dict) or not detailed.get("form"):
        raise AutomationError("API не вернул структуру шаблона существующего проекта")
    return detailed


def payload_from_existing_project(
    existing: dict[str, Any],
    template: dict[str, Any],
    project: ProjectData,
    region_record: dict[str, Any],
) -> dict[str, Any]:
    values = values_from_project_response(existing) or template_default_values(template)
    project_type = existing.get("type") or {}
    existing_region = existing.get("region") or {}
    return {
        "typeID": project_type.get("ID") or template.get("typeID"),
        "templateID": existing.get("templateID") or template.get("ID"),
        "values": values,
        # Название и регион приводим к данным Word, остальные поля сохраняем.
        "name": project.title,
        "isDraft": True,
        "regionFiasID": region_record.get("fiasID") or existing_region.get("fiasID"),
        "costs": copy.deepcopy(existing.get("costs") or []),
        "team": copy.deepcopy(existing.get("team") or []),
    }


def _labels_equal(value: str, aliases: Sequence[str]) -> bool:
    normalized = normalize_text(value)
    for alias in aliases:
        target = normalize_text(alias)
        if normalized == target:
            return True
    return False


def _find_tab(template: dict[str, Any], aliases: Sequence[str]) -> dict[str, Any] | None:
    for tab in (template.get("form") or {}).get("tabs") or []:
        if _labels_equal(str(tab.get("label", "")), aliases):
            return tab
    return None


def _find_group(
    tab: dict[str, Any] | None,
    aliases: Sequence[str],
) -> dict[str, Any] | None:
    if not tab:
        return None
    for group in tab.get("groups") or []:
        if _labels_equal(str(group.get("label", "")), aliases):
            return group
    return None


def _find_field(
    group: dict[str, Any] | None,
    aliases: Sequence[str],
) -> dict[str, Any] | None:
    if not group:
        return None
    for field_info in group.get("fields") or []:
        if _labels_equal(str(field_info.get("label", "")), aliases):
            return field_info
    return None


def _find_part(
    field_info: dict[str, Any] | None,
    aliases: Sequence[str],
) -> dict[str, Any] | None:
    if not field_info:
        return None
    for part in field_info.get("parts") or []:
        if _labels_equal(str(part.get("label", "")), aliases):
            return part
    return None


def _default_composite_row(field_info: dict[str, Any]) -> dict[str, Any]:
    row: dict[str, Any] = {}
    for part in field_info.get("parts") or []:
        if not isinstance(part, dict):
            continue
        if part.get("type") == "composite":
            row[str(part["ID"])] = [_default_composite_row(part)]
        else:
            row[str(part["ID"])] = copy.deepcopy(part.get("value"))
    return row


def template_default_values(template: dict[str, Any]) -> dict[str, Any]:
    values: dict[str, Any] = {}
    for tab in (template.get("form") or {}).get("tabs") or []:
        for group in tab.get("groups") or []:
            for field_info in group.get("fields") or []:
                field_id = str(field_info.get("ID", ""))
                if not field_id:
                    continue
                value = copy.deepcopy(field_info.get("value"))
                if field_info.get("type") == "composite" and value is None:
                    value = [_default_composite_row(field_info)]
                values[field_id] = value
    return values


def values_from_project_response(project_response: dict[str, Any]) -> dict[str, Any]:
    form = project_response.get("form") or {}
    values: dict[str, Any] = {}
    for tab in form.get("tabs") or []:
        for group in tab.get("groups") or []:
            for field_info in group.get("fields") or []:
                field_id = str(field_info.get("ID", ""))
                if field_id:
                    values[field_id] = copy.deepcopy(field_info.get("value"))
    return values


def _canonical_option(field_info: dict[str, Any] | None, value: str) -> str:
    value = clean_text(value)
    if not value:
        return ""
    items = field_info.get("items") if field_info else None
    if isinstance(items, list):
        for item in items:
            if normalize_text(str(item)) == normalize_text(value):
                return str(item)
        scored = sorted(
            (
                (project_score(value, str(item)), str(item))
                for item in items
                if clean_text(str(item))
            ),
            reverse=True,
        )
        if scored and scored[0][0] >= 0.75:
            return scored[0][1]
    return value


def _parse_number_value(value: str, default: int | float | None = None) -> int | float | None:
    text = clean_text(value).replace("\u00a0", " ").replace(" ", "").replace(",", ".")
    match = re.search(r"-?\d+(?:\.\d+)?", text)
    if not match:
        return default
    try:
        number = float(match.group(0))
    except ValueError:
        return default
    return int(number) if number.is_integer() else number


def _api_date(value: str, *, end_of_month: bool = False) -> str | None:
    text = clean_text(value)
    if not text:
        return None

    patterns = (
        (r"\b(\d{2})\.(\d{2})\.(\d{4})\b", "day"),
        (r"\b(\d{4})-(\d{2})-(\d{2})\b", "iso"),
        (r"\b(\d{2})\.(\d{4})\b", "month"),
        (r"\b(\d{4})-(\d{2})\b", "iso_month"),
    )
    for pattern, kind in patterns:
        match = re.search(pattern, text)
        if not match:
            continue
        try:
            if kind == "day":
                day, month, year = map(int, match.groups())
            elif kind == "iso":
                year, month, day = map(int, match.groups())
            elif kind == "month":
                month, year = map(int, match.groups())
                day = calendar.monthrange(year, month)[1] if end_of_month else 1
            else:
                year, month = map(int, match.groups())
                day = calendar.monthrange(year, month)[1] if end_of_month else 1
            parsed = date(year, month, day)
            return parsed.strftime("%Y-%m-%d 00:00:00")
        except ValueError:
            continue
    return None


def _api_date_interval(start: str, end: str) -> str | None:
    left = _api_date(start, end_of_month=False)
    right = _api_date(end, end_of_month=True)
    if not left and right:
        left = right
    if not right and left:
        right = left
    if not left or not right:
        return None
    return f"{left} - {right}"


def _latest_project_date(project: ProjectData) -> str | None:
    candidates: list[date] = []
    for raw in [project.end_date, *(event.deadline for event in project.calendar)]:
        formatted = _api_date(raw, end_of_month=True)
        if not formatted:
            continue
        try:
            candidates.append(datetime.strptime(formatted, "%Y-%m-%d %H:%M:%S").date())
        except ValueError:
            continue
    if not candidates:
        return None
    return max(candidates).strftime("%Y-%m-%d 00:00:00")


async def resolve_region_record(
    context: BrowserContext,
    settings: Settings,
    authorization: str,
    region_name: str,
) -> dict[str, Any]:
    region_name = clean_text(region_name)
    if not region_name:
        raise AutomationError("В Word-файле отсутствует регион реализации проекта")

    result = await api_json_request(
        context,
        settings,
        authorization,
        "GET",
        f"/geo/regions?search={quote_plus(region_name)}&active=1",
    )
    if not isinstance(result, list):
        result = []

    exact = [
        item for item in result
        if normalize_text(str(item.get("name", ""))) == normalize_text(region_name)
    ]
    if exact:
        return exact[0]

    scored = sorted(
        (
            (project_score(region_name, str(item.get("name", ""))), item)
            for item in result
            if item.get("fiasID")
        ),
        key=lambda pair: pair[0],
        reverse=True,
    )
    if scored and scored[0][0] >= 0.62:
        return scored[0][1]

    all_regions = await api_json_request(
        context,
        settings,
        authorization,
        "GET",
        "/open/geo/regions?per-page=100&active=1",
    )
    scored = sorted(
        (
            (project_score(region_name, str(item.get("name", ""))), item)
            for item in (all_regions if isinstance(all_regions, list) else [])
            if item.get("fiasID")
        ),
        key=lambda pair: pair[0],
        reverse=True,
    )
    if scored and scored[0][0] >= 0.62:
        return scored[0][1]

    raise AutomationError(f"Регион «{region_name}» не найден в справочнике ФГАИС")


async def get_project_template_api(
    context: BrowserContext,
    settings: Settings,
    authorization: str,
) -> dict[str, Any]:
    templates = await api_json_request(
        context,
        settings,
        authorization,
        "GET",
        "/project-templates?isProgram=0&templateIsParent=1&per-page=100",
    )
    if not isinstance(templates, list):
        raise AutomationError("API вернул некорректный список шаблонов")

    expected = normalize_text(settings.project_template_name)
    matches = [
        item for item in templates
        if normalize_text(str(item.get("displayName") or item.get("name") or "")) == expected
    ]
    if not matches:
        matches = sorted(
            (
                item for item in templates
                if project_score(
                    settings.project_template_name,
                    str(item.get("displayName") or item.get("name") or ""),
                ) >= 0.82
            ),
            key=lambda item: project_score(
                settings.project_template_name,
                str(item.get("displayName") or item.get("name") or ""),
            ),
            reverse=True,
        )
    if not matches:
        raise AutomationError(
            "Через API не найден шаблон: " + settings.project_template_name
        )

    template_id = str(matches[0].get("ID", ""))
    if not template_id:
        raise AutomationError("У найденного шаблона отсутствует ID")
    return await get_project_template_by_id_api(
        context,
        settings,
        authorization,
        template_id,
    )


def _set_field_value(
    values: dict[str, Any],
    field_info: dict[str, Any] | None,
    value: Any,
) -> int:
    if not field_info or value in (None, "", []):
        return 0
    values[str(field_info["ID"])] = value
    return 1


def _append_missing(skipped: list[str], label: str) -> None:
    skipped.append(f"В Word отсутствует: {label}")


def apply_general_api_stage(
    values: dict[str, Any],
    template: dict[str, Any],
    project: ProjectData,
) -> tuple[int, list[str]]:
    skipped: list[str] = []
    count = 0
    tab = _find_tab(template, ("Общее",))
    group = _find_group(tab, ("Общая информация",))

    scale_field = _find_field(
        group, ("Масштаб реализации проекта", "Масштаб проекта")
    )
    if project.scale:
        count += _set_field_value(
            values, scale_field, _canonical_option(scale_field, project.scale)
        )
    else:
        _append_missing(skipped, "Масштаб реализации проекта")

    date_field = _find_field(
        group, ("Дата начала и окончания проекта", "Сроки реализации проекта")
    )
    date_value = _api_date_interval(project.start_date, project.end_date)
    if date_value:
        count += _set_field_value(values, date_field, date_value)
    else:
        _append_missing(skipped, "Дата начала и окончания проекта")

    skipped.extend(
        [
            "Логотип не переносится из Word",
            "Личные поля автора, резюме и видеовизитка не переносятся",
        ]
    )
    return count, skipped


async def apply_about_api_stage(
    values: dict[str, Any],
    template: dict[str, Any],
    project: ProjectData,
    context: BrowserContext,
    settings: Settings,
    authorization: str,
    region_cache: dict[str, dict[str, Any]],
) -> tuple[int, list[str]]:
    skipped: list[str] = []
    count = 0
    tab = _find_tab(template, ("О проекте",))
    info_group = _find_group(tab, ("Информация о проекте",))

    fields = (
        (("Краткая информация о проекте",), project.summary),
        (("Основные целевые группы, на которые направлен проект",), project.target_groups),
        (
            (
                "Описание проблемы, решению/снижению которой посвящен проект",
                "Описание проблемы, решению которой посвящен проект",
            ),
            project.problem,
        ),
        (("Основная цель проекта",), project.goal),
        (
            (
                "Опыт команды проекта по реализации социально значимых проектов и/или событий",
            ),
            project.team_experience,
        ),
        (("Перспектива развития и потенциал проекта",), project.development),
    )
    for aliases, value in fields:
        field_info = _find_field(info_group, aliases)
        if value:
            count += _set_field_value(values, field_info, value)
        else:
            _append_missing(skipped, aliases[0])

    task_group = _find_group(tab, ("Задачи",))
    task_composite = _find_field(task_group, ("Добавить задачу",))
    task_part = _find_part(task_composite, ("Поставленная задача", "Задача"))
    unique_tasks: list[str] = []
    seen_tasks: set[str] = set()
    for event in project.calendar:
        task = clean_text(event.task)
        key = normalize_text(task)
        if task and key not in seen_tasks:
            seen_tasks.add(key)
            unique_tasks.append(task)
    if task_composite and task_part and unique_tasks:
        values[str(task_composite["ID"])] = [
            {str(task_part["ID"]): task} for task in unique_tasks
        ]
        count += len(unique_tasks)
    elif not unique_tasks:
        skipped.append("Задачи не распознаны из календарного плана")

    geography_group = _find_group(tab, ("География проекта",))
    geography_composite = _find_field(geography_group, ("Поле", "Добавить поле"))
    geography_part = _find_part(
        geography_composite, ("Выберите регион или федеральный округ",)
    )
    address_part = _find_part(geography_composite, ("Адрес",))
    locations = project.geography[:] or ([project.region] if project.region else [])
    geography_rows: list[dict[str, Any]] = []
    for location in locations:
        location = clean_text(location)
        search_region = location if "республик" in normalize_text(location) or "област" in normalize_text(location) or "край" in normalize_text(location) else project.region
        cache_key = normalize_text(search_region)
        try:
            region_record = region_cache.get(cache_key)
            if region_record is None:
                region_record = await resolve_region_record(
                    context, settings, authorization, search_region
                )
                region_cache[cache_key] = region_record
            row: dict[str, Any] = {}
            if geography_part:
                row[str(geography_part["ID"])] = region_record.get("fiasID")
            if address_part:
                region_label = clean_text(str(region_record.get("name", "")))
                row[str(address_part["ID"])] = (
                    None
                    if normalize_text(location) == normalize_text(region_label)
                    else location
                )
            geography_rows.append(row)
        except Exception as exc:
            skipped.append(f"География «{location}» пропущена: {exc}")

    if geography_composite and geography_rows:
        values[str(geography_composite["ID"])] = geography_rows
        count += len(geography_rows)
    elif not locations:
        _append_missing(skipped, "География проекта")

    skipped.extend(
        [
            "Документы, подтверждающие проблему, не восстанавливаются из Word",
            "Регион нахождения автора проекта пропущен как персональный блок",
        ]
    )
    return count, skipped


def apply_results_api_stage(
    values: dict[str, Any],
    template: dict[str, Any],
    project: ProjectData,
) -> tuple[int, list[str]]:
    skipped: list[str] = []
    count = 0
    tab = _find_tab(template, ("Результаты",))

    latest_date = _latest_project_date(project)
    date_group = _find_group(
        tab,
        (
            "Дата плановых значений результатов Предоставления субсидии по годам (срокам) реализации Соглашения",
        ),
    )
    date_field = _find_field(
        date_group,
        (
            "Дата плановых значений результатов Предоставления субсидии по годам (срокам) реализации Соглашения",
        ),
    )
    count += _set_field_value(values, date_field, latest_date)

    results: tuple[tuple[Sequence[str], int | float | None, str], ...] = (
        (
            ("Количество мероприятий, проведенных в рамках проекта",),
            _parse_number_value(project.events_count, len(project.calendar) or None),
            "Ед.",
        ),
        (
            ("Количество участников мероприятий, вовлеченных в реализацию проекта",),
            _parse_number_value(
                project.participants_count,
                sum(
                    int(_parse_number_value(event.unique_participants, 0) or 0)
                    for event in project.calendar
                ) or None,
            ),
            "Чел.",
        ),
        (
            (
                "Количество публикаций о мероприятиях проекта в средствах массовой информации, а также в информационно-телекоммуникационной сети «Интернет»",
            ),
            _parse_number_value(
                project.publications_count,
                sum(
                    int(_parse_number_value(event.publications, 0) or 0)
                    for event in project.calendar
                ) or None,
            ),
            "Ед.",
        ),
        (
            (
                "Количество просмотров публикаций о мероприятиях проекта в информационно-телекоммуникационной сети «Интернет»",
            ),
            _parse_number_value(
                project.views_count,
                sum(
                    int(_parse_number_value(event.views, 0) or 0)
                    for event in project.calendar
                ) or None,
            ),
            "Ед.",
        ),
    )
    for group_aliases, planned, unit in results:
        group = _find_group(tab, group_aliases)
        amount_field = _find_field(group, ("Плановое количество",))
        unit_field = _find_field(group, ("Ед. измерения",))
        if planned is not None:
            count += _set_field_value(values, amount_field, planned)
            count += _set_field_value(
                values, unit_field, _canonical_option(unit_field, unit)
            )
        else:
            _append_missing(skipped, group_aliases[0])
        if group_aliases[0].startswith("Количество мероприятий"):
            deadline_field = _find_field(group, ("Крайняя дата проведения",))
            count += _set_field_value(values, deadline_field, latest_date)
        if group_aliases[0].startswith("Количество просмотров"):
            social_field = _find_field(group, ("Социальный эффект",))
            if project.social_effect:
                count += _set_field_value(values, social_field, project.social_effect)
            else:
                _append_missing(skipped, "Социальный эффект")

    return count, skipped


def apply_calendar_api_stage(
    values: dict[str, Any],
    template: dict[str, Any],
    project: ProjectData,
) -> tuple[int, list[str]]:
    skipped: list[str] = []
    count = 0
    tab = _find_tab(template, ("Календарный план",))
    group = _find_group(tab, ("Задачи",))
    tasks_field = _find_field(group, ("Добавить задачу",))
    task_part = _find_part(tasks_field, ("Поставленная задача", "Задача"))
    events_part = _find_part(tasks_field, ("Добавить мероприятие",))
    if not tasks_field or not task_part or not events_part:
        raise AutomationError("В шаблоне API не найдена структура календарного плана")

    event_parts = {
        normalize_text(str(part.get("label", ""))): part
        for part in events_part.get("parts") or []
        if isinstance(part, dict)
    }

    grouped: dict[str, tuple[str, list[CalendarEvent]]] = {}
    order: list[str] = []
    for event in project.calendar:
        task = clean_text(event.task) or "Реализовать комплекс мероприятий проекта"
        key = normalize_text(task)
        if key not in grouped:
            grouped[key] = (task, [])
            order.append(key)
        grouped[key][1].append(event)

    if not grouped:
        skipped.append("Календарный план не распознан в Word")
        return 0, skipped

    task_rows: list[dict[str, Any]] = []
    for key in order:
        task_text, events = grouped[key]
        event_rows: list[dict[str, Any]] = []
        for event in events:
            row: dict[str, Any] = {}
            mapping: tuple[tuple[str, Any], ...] = (
                ("название мероприятия", clean_text(event.title) or None),
                ("крайняя дата выполнения", _api_date(event.deadline, end_of_month=True)),
                ("описание мероприятия", clean_text(event.description) or None),
                (
                    "количество уникальных участников",
                    _parse_number_value(event.unique_participants, 0),
                ),
                (
                    "количество повторяющихся участников",
                    _parse_number_value(event.repeat_participants, 0),
                ),
                ("количество публикаций", _parse_number_value(event.publications, 0)),
                ("количество просмотров", _parse_number_value(event.views, 0)),
                ("дополнительная информация", clean_text(event.additional) or None),
            )
            for label, value in mapping:
                part = event_parts.get(normalize_text(label))
                if part:
                    row[str(part["ID"])] = value
                    if value not in (None, "", 0):
                        count += 1
            event_rows.append(row)
            if not event.title:
                skipped.append("У мероприятия отсутствует название")
            if not event.deadline:
                skipped.append(
                    f"У мероприятия «{event.title or 'без названия'}» отсутствует дата"
                )
        task_rows.append(
            {
                str(task_part["ID"]): task_text,
                str(events_part["ID"]): event_rows,
            }
        )
        count += 1

    values[str(tasks_field["ID"])] = task_rows
    return count, skipped


def _support_types(field_info: dict[str, Any] | None, value: str) -> list[str]:
    raw = [
        clean_text(item)
        for item in re.split(r"[,;\n/]+", clean_text(value))
        if clean_text(item)
    ]
    result: list[str] = []
    for item in raw:
        canonical = _canonical_option(field_info, item)
        if canonical and canonical not in result:
            result.append(canonical)
    return result


def apply_cofinancing_api_stage(
    values: dict[str, Any],
    template: dict[str, Any],
    project: ProjectData,
) -> tuple[int, list[str]]:
    skipped: list[str] = [
        "Подтверждающие файлы софинансирования не восстанавливаются из Word"
    ]
    count = 0
    tab = _find_tab(template, ("Софинансирование",))

    own_group = _find_group(tab, ("Собственные средства",))
    own_expenses_field = _find_field(own_group, ("Перечень расходов",))
    own_amount_composite = _find_field(own_group, ("Поле",))
    own_amount_part = _find_part(own_amount_composite, ("Сумма, руб.", "Сумма"))
    own_entries = [entry for entry in project.cofinancing if entry.kind == "own"]
    if own_entries:
        descriptions = "\n".join(
            clean_text(entry.expenses) for entry in own_entries if clean_text(entry.expenses)
        )
        total = sum(
            float(_parse_number_value(entry.amount, 0) or 0)
            for entry in own_entries
        )
        count += _set_field_value(values, own_expenses_field, descriptions)
        if own_amount_composite and own_amount_part and total > 0:
            values[str(own_amount_composite["ID"])] = [
                {
                    str(own_amount_part["ID"]): int(total) if total.is_integer() else total,
                    **{
                        str(part["ID"]): None
                        for part in own_amount_composite.get("parts") or []
                        if isinstance(part, dict) and part.get("ID") != own_amount_part.get("ID")
                    },
                }
            ]
            count += 1

    partner_group = _find_group(tab, ("Партнер", "Партнёр"))
    partner_composite = _find_field(partner_group, ("Партнера", "Партнёра", "Партнер"))
    support_part = _find_part(partner_composite, ("Тип поддержки",))
    file_part = _find_part(
        partner_composite,
        ("Файл, подтверждающий готовность предоставить партнёрскую поддержку",),
    )
    name_part = _find_part(partner_composite, ("Название партнера", "Название партнёра"))
    expenses_part = _find_part(partner_composite, ("Перечень расходов",))
    amount_part = _find_part(partner_composite, ("Сумма, руб.", "Сумма"))

    rows: list[dict[str, Any]] = []
    for entry in [item for item in project.cofinancing if item.kind == "partner"]:
        if not any(
            clean_text(value)
            for value in (
                entry.partner_name,
                entry.support_type,
                entry.expenses,
                entry.amount,
            )
        ):
            continue
        row: dict[str, Any] = {}
        if support_part:
            row[str(support_part["ID"])] = _support_types(
                support_part, entry.support_type
            )
        if file_part:
            row[str(file_part["ID"])] = None
        if name_part:
            row[str(name_part["ID"])] = clean_text(entry.partner_name) or None
        if expenses_part:
            row[str(expenses_part["ID"])] = clean_text(entry.expenses) or None
        if amount_part:
            row[str(amount_part["ID"])] = _parse_number_value(entry.amount)
        rows.append(row)
        count += sum(value not in (None, "", [], 0) for value in row.values())

    if partner_composite and rows:
        values[str(partner_composite["ID"])] = rows
    return count, skipped


def _cost_category_for_item(
    item: ExpenseItem,
    categories: list[dict[str, Any]],
) -> dict[str, Any] | None:
    category_text = clean_text(item.category)
    normalized = normalize_text(category_text)

    # Порядок категорий Word-выгрузки совпадает с порядком категорий API.
    for index, source_category in enumerate(EXPENSE_CATEGORIES):
        if normalize_text(source_category) == normalized and index < len(categories):
            return categories[index]

    scored = sorted(
        (
            (
                max(
                    project_score(category_text, str(category.get("name", ""))),
                    project_score(
                        category_text,
                        EXPENSE_CATEGORIES[index] if index < len(EXPENSE_CATEGORIES) else "",
                    ),
                ),
                category,
            )
            for index, category in enumerate(categories)
        ),
        key=lambda pair: pair[0],
        reverse=True,
    )
    if scored and scored[0][0] >= 0.48:
        return scored[0][1]
    return None


def _cost_type(category: dict[str, Any], item_name: str) -> str:
    allow_goods = bool(category.get("allowGoods"))
    allow_service = bool(category.get("allowService"))
    if allow_goods and not allow_service:
        return "goods"
    if allow_service and not allow_goods:
        return "service"

    service_markers = (
        "услуг",
        "аренд",
        "разработ",
        "изготов",
        "печать",
        "достав",
        "перевоз",
        "прожив",
        "питани",
        "поддержк",
        "размещени",
    )
    normalized = normalize_text(item_name)
    return "service" if any(marker in normalized for marker in service_markers) else "goods"


def build_costs_api(
    template: dict[str, Any],
    project: ProjectData,
) -> tuple[list[dict[str, Any]], int, list[str]]:
    categories = [
        item for item in template.get("costCategories") or [] if isinstance(item, dict)
    ]
    costs: list[dict[str, Any]] = []
    skipped: list[str] = []
    count = 0

    for item in project.expenses:
        name = clean_text(item.item)
        price = _parse_number_value(item.price)
        quantity = _parse_number_value(item.quantity, 1)
        if not name:
            skipped.append("Строка сметы без наименования пропущена")
            continue
        if price is None or float(price) <= 0:
            skipped.append(f"Позиция «{name}» пропущена: некорректная цена")
            continue
        if quantity is None or float(quantity) <= 0:
            skipped.append(f"Позиция «{name}» пропущена: некорректное количество")
            continue

        category = _cost_category_for_item(item, categories)
        if not category:
            skipped.append(
                f"Позиция «{name}» пропущена: не определена категория расходов"
            )
            continue

        costs.append(
            {
                "name": name,
                "description": clean_text(item.justification) or None,
                "price": f"{float(price):.2f}",
                "count": int(quantity) if float(quantity).is_integer() else quantity,
                "type": _cost_type(category, name),
                "categoryID": category.get("ID"),
            }
        )
        count += 1
    return costs, count, skipped


async def save_project_api(
    context: BrowserContext,
    settings: Settings,
    authorization: str,
    project_id: str,
    payload: dict[str, Any],
    labels: dict[str, str],
) -> dict[str, Any]:
    result = await api_json_request(
        context,
        settings,
        authorization,
        "POST",
        f"/projects/{project_id}"
        "?expand=agreement,agreementStatusCode,form,costs,won,wonAmount,"
        "programInfo,isAllFieldsDisabled",
        payload=payload,
        labels=labels,
    )
    if not isinstance(result, dict):
        raise AutomationError("API сохранения проекта не вернул объект проекта")
    return result


async def run_api_stage(
    stage_name: str,
    current_payload: dict[str, Any],
    builder,
    saver,
    skipped: list[str],
) -> tuple[dict[str, Any], int]:
    while True:
        candidate = copy.deepcopy(current_payload)
        try:
            count, stage_skipped = await builder(candidate)
            await saver(candidate)
            skipped.extend(stage_skipped)
            print(f"Этап «{stage_name}»: сохранено полей/строк — {count}")
            return candidate, count
        except KeyboardInterrupt:
            raise
        except Exception as exc:
            action = await ask_error_action(
                stage_name,
                exc,
                skip_label="пропустить только этот раздел и продолжить этот же аккаунт",
            )
            if action == "retry":
                continue
            if action == "skip":
                skipped.append(f"Раздел «{stage_name}» пропущен: {exc}")
                return current_payload, 0
            raise KeyboardInterrupt


async def process_upload_account_api(
    browser: Browser,
    account: TargetAccount,
    settings: Settings,
) -> TransferResult:
    result = TransferResult(
        row_number=account.row_number,
        fio=account.fio,
        operation="создание черновика через API",
    )
    try:
        word_file = resolve_word_file(account.word_reference, settings)
        project = parse_project_docx(word_file, account.title_override)
        result.word_file = word_file.name
        result.project_name = project.title
        result.warnings.extend(project.parser_warnings)
    except Exception as exc:
        result.status = "ошибка"
        result.error = str(exc)
        return result

    context: BrowserContext = await browser.new_context(
        viewport={"width": 1360, "height": 900},
        accept_downloads=False,
    )
    context.set_default_timeout(settings.timeout_ms)
    if settings.block_heavy_resources:
        async def route_handler(route):
            if route.request.resource_type in {"image", "media", "font"}:
                await route.abort()
            else:
                await route.continue_()
        await context.route("**/*", route_handler)

    page = await context.new_page()
    auth_capture: dict[str, str] = {}

    def capture_request(request) -> None:
        if request.url.startswith(f"{settings.base_url}/api/"):
            authorization = request.headers.get("authorization", "")
            if authorization:
                auth_capture["authorization"] = authorization

    page.on("request", capture_request)
    skipped: list[str] = []
    target_project_id = ""
    new_project_created = False
    existing_project_selected = False

    try:
        login_ok, _ = await run_optional_website_step(
            "авторизация в аккаунте",
            lambda: login_account(page, account.login, account.password, settings),
            skipped,
        )
        if not login_ok:
            raise AutomationError(
                "Авторизация была пропущена; API-перенос без авторизации невозможен"
            )

        authorization = await resolve_authorization_header(
            page, auth_capture, settings
        )

        # До создания нового черновика проверяем точное совпадение итогового названия.
        while True:
            try:
                account_projects = await get_account_projects_api(
                    context, settings, authorization
                )
                destination_mode, selected_existing = await choose_existing_project_or_new(
                    account_projects, project.title
                )
                break
            except KeyboardInterrupt:
                raise
            except Exception as exc:
                action = await ask_error_action(
                    "проверка проектов с таким же названием",
                    exc,
                    skip_label="пропустить проверку и создать новый проект",
                )
                if action == "retry":
                    continue
                if action == "skip":
                    destination_mode, selected_existing = "new", None
                    skipped.append(
                        f"Проверка совпадающих проектов пропущена: {exc}"
                    )
                    break
                raise KeyboardInterrupt

        if destination_mode == "existing":
            if not selected_existing or not _existing_project_is_editable(selected_existing):
                raise AutomationError(
                    "Выбранный существующий проект не является доступным черновиком"
                )

            target_project_id = clean_text(str(selected_existing.get("ID", "")))
            existing_detail = await get_project_detail_api(
                context,
                settings,
                authorization,
                target_project_id,
            )
            if not bool(existing_detail.get("isDraft", selected_existing.get("isDraft"))):
                raise AutomationError(
                    "Существующий проект уже не является черновиком и не может быть перезаписан"
                )
            if not bool(
                existing_detail.get(
                    "isManageable", selected_existing.get("isManageable")
                )
            ):
                raise AutomationError(
                    "У текущего аккаунта нет права редактирования выбранного проекта"
                )

            template_id = clean_text(str(existing_detail.get("templateID", "")))
            if not template_id:
                existing_template = existing_detail.get("template") or {}
                template_id = clean_text(str(existing_template.get("ID", "")))
            template = await get_project_template_by_id_api(
                context,
                settings,
                authorization,
                template_id,
            )
            labels = _schema_label_map(template)
            region_record = await resolve_region_record(
                context, settings, authorization, project.region
            )
            current_payload = payload_from_existing_project(
                existing_detail,
                template,
                project,
                region_record,
            )
            existing_project_selected = True
            result.operation = "заполнение существующего черновика через API"
            result.warnings.append(
                f"Заполнен существующий черновик ID {target_project_id}; "
                "новый проект не создавался"
            )
            total_filled = 2  # название и регион актуализированы по Word
        else:
            template = await get_project_template_api(
                context, settings, authorization
            )
            labels = _schema_label_map(template)
            region_record = await resolve_region_record(
                context, settings, authorization, project.region
            )

            initial_values = template_default_values(template)
            initial_payload: dict[str, Any] = {
                "typeID": template.get("typeID"),
                "templateID": template.get("ID"),
                "values": initial_values,
                "name": project.title,
                "isDraft": True,
                "regionFiasID": region_record.get("fiasID"),
                "costs": [],
                "team": None,
            }

            while True:
                try:
                    created = await api_json_request(
                        context,
                        settings,
                        authorization,
                        "POST",
                        "/projects",
                        payload=initial_payload,
                        labels=labels,
                    )
                    if not isinstance(created, dict) or not created.get("ID"):
                        raise AutomationError(
                            "API создания черновика не вернул ID проекта"
                        )
                    target_project_id = str(created["ID"])
                    new_project_created = True
                    break
                except Exception as exc:
                    action = await ask_error_action(
                        "создание пустого черновика",
                        exc,
                        skip_label=(
                            "пропустить создание черновика "
                            "(остальные разделы этого проекта сохранить будет невозможно)"
                        ),
                    )
                    if action == "retry":
                        continue
                    if action == "skip":
                        skipped.append(f"Создание черновика пропущено: {exc}")
                        result.status = "частично"
                        result.fields_skipped = skipped
                        return result
                    raise KeyboardInterrupt

            initial_values = values_from_project_response(created) or initial_values
            current_payload = {
                **initial_payload,
                "values": initial_values,
                "costs": [],
                "team": [],
            }
            total_filled = 2  # название и регион

        region_cache = {normalize_text(project.region): region_record}

        async def saver(payload: dict[str, Any]) -> None:
            await save_project_api(
                context,
                settings,
                authorization,
                target_project_id,
                payload,
                labels,
            )

        async def general_builder(payload: dict[str, Any]):
            count, stage_skipped = apply_general_api_stage(
                payload["values"], template, project
            )
            return count, stage_skipped

        async def about_builder(payload: dict[str, Any]):
            return await apply_about_api_stage(
                payload["values"],
                template,
                project,
                context,
                settings,
                authorization,
                region_cache,
            )

        async def results_builder(payload: dict[str, Any]):
            return apply_results_api_stage(
                payload["values"], template, project
            )

        async def calendar_builder(payload: dict[str, Any]):
            return apply_calendar_api_stage(
                payload["values"], template, project
            )

        async def cofinancing_builder(payload: dict[str, Any]):
            return apply_cofinancing_api_stage(
                payload["values"], template, project
            )

        async def expenses_builder(payload: dict[str, Any]):
            costs, count, stage_skipped = build_costs_api(template, project)
            payload["costs"] = costs
            return count, stage_skipped

        stages = (
            ("Общее", general_builder),
            ("О проекте", about_builder),
            ("Результаты", results_builder),
            ("Календарный план", calendar_builder),
            ("Софинансирование", cofinancing_builder),
            ("Расходы", expenses_builder),
        )
        for stage_name, builder in stages:
            current_payload, stage_count = await run_api_stage(
                stage_name,
                current_payload,
                builder,
                saver,
                skipped,
            )
            total_filled += stage_count

        result.fields_filled = total_filled
        result.fields_skipped = skipped + [
            "Вкладка «Команда» пропущена по заданию",
            "Вкладка «Медиа» пропущена по заданию",
            "Бинарные вложения из Word не восстанавливаются",
        ]
        result.status = "успешно" if total_filled > 2 else "частично"

        edit_url = (
            f"{settings.base_url}/projects/edit/{target_project_id}"
            "?backTo=projects&noreset=true"
        )
        try:
            await page.goto(
                edit_url,
                wait_until="domcontentloaded",
                timeout=max(settings.timeout_ms, 30_000),
            )
            await dismiss_cookie_banner(page)
            await page.wait_for_timeout(1200)
        except Exception as exc:
            result.warnings.append(
                "Проект сохранён, но итоговая страница не открылась: "
                f"{exc}"
            )

        await take_screenshot(
            page,
            settings,
            account.row_number,
            account.fio,
            (
                "СУЩЕСТВУЮЩИЙ_ЧЕРНОВИК_ЗАПОЛНЕН_API"
                if existing_project_selected
                else "ЧЕРНОВИК_СОЗДАН_API"
            ),
        )
        return result
    except KeyboardInterrupt:
        raise
    except Exception as exc:
        result.status = "ошибка"
        result.error = str(exc)
        result.fields_skipped = skipped
        if target_project_id:
            if new_project_created:
                result.warnings.append(
                    f"Пустой или частично заполненный новый черновик создан: "
                    f"{target_project_id}"
                )
            elif existing_project_selected:
                result.warnings.append(
                    f"Существующий черновик мог быть частично обновлён: "
                    f"{target_project_id}"
                )
        await take_screenshot(
            page,
            settings,
            account.row_number,
            account.fio,
            "ОШИБКА_API_ПЕРЕНОСА",
        )
        return result
    finally:
        await context.close()


async def process_upload_account_legacy_dom(
    browser: Browser,
    account: TargetAccount,
    settings: Settings,
) -> TransferResult:
    result = TransferResult(
        row_number=account.row_number,
        fio=account.fio,
        operation="создание черновика",
    )
    try:
        word_file = resolve_word_file(account.word_reference, settings)
        project = parse_project_docx(word_file, account.title_override)
        result.word_file = word_file.name
        result.project_name = project.title
        result.warnings.extend(project.parser_warnings)
    except Exception as exc:
        result.status = "ошибка"
        result.error = str(exc)
        return result

    context: BrowserContext = await browser.new_context(
        viewport={"width": 1360, "height": 900}, accept_downloads=False
    )
    context.set_default_timeout(settings.timeout_ms)
    if settings.block_heavy_resources:
        async def route_handler(route):
            if route.request.resource_type in {"image", "media", "font"}:
                await route.abort()
            else:
                await route.continue_()
        await context.route("**/*", route_handler)
    page = await context.new_page()

    skipped_steps: list[str] = []
    try:
        await run_optional_website_step(
            "авторизация в аккаунте",
            lambda: login_account(page, account.login, account.password, settings),
            skipped_steps,
        )

        template_ok, _ = await run_optional_website_step(
            "открытие шаблона нового проекта",
            lambda: open_new_project_template(page, settings),
            skipped_steps,
        )

        initial_filled = 0
        identity_filled = False
        draft_created = False
        if template_ok:
            initial_filled, identity_filled, draft_created = await create_initial_draft(
                page,
                project,
                settings,
                skipped_steps,
            )
        else:
            skipped_steps.append(
                "Стартовые поля не заполнялись, потому что шаг открытия шаблона был пропущен"
            )

        if not draft_created:
            skipped_steps.append(
                "Режим черновика не подтверждён; последующие вкладки будут обрабатываться по возможности"
            )

        filled, skipped = await fill_project(
            page,
            project,
            settings,
            identity_already_filled=identity_filled,
        )
        result.fields_filled = initial_filled + filled
        result.fields_skipped = skipped_steps + skipped
        if draft_created and result.fields_filled:
            result.status = "успешно"
        elif result.fields_filled:
            result.status = "частично"
        else:
            result.status = "частично"
        await take_screenshot(
            page, settings, account.row_number, account.fio, "ЧЕРНОВИК_ОБРАБОТАН"
        )
        return result
    except KeyboardInterrupt:
        raise
    except Exception as exc:
        result.status = "ошибка"
        result.error = str(exc)
        result.fields_skipped = skipped_steps
        await take_screenshot(
            page, settings, account.row_number, account.fio, "ОШИБКА_ПЕРЕНОСА"
        )
        return result
    finally:
        await context.close()


async def process_upload_account(
    browser: Browser,
    account: TargetAccount,
    settings: Settings,
) -> TransferResult:
    """Основной режим: создание и заполнение черновика через фактический API сайта."""
    return await process_upload_account_api(browser, account, settings)


# ---------------------------------------------------------------------------
# Отчёты и запуск
# ---------------------------------------------------------------------------


def write_reports(results: list[TransferResult], settings: Settings) -> tuple[Path, Path]:
    settings.logs_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    csv_path = settings.logs_dir / f"transfer_report_{stamp}.csv"
    xls_path = settings.logs_dir / f"transfer_report_{stamp}.xls"
    headers = [
        "Строка XLS",
        "ФИО",
        "Операция",
        "Проект",
        "Word-файл",
        "Заполнено полей",
        "Пропущено / не переносится",
        "Предупреждения",
        "Статус",
        "Ошибка",
    ]

    rows = [
        [
            item.row_number,
            item.fio,
            item.operation,
            item.project_name,
            item.word_file,
            item.fields_filled,
            item.skipped_text(),
            item.warnings_text(),
            item.status,
            item.error,
        ]
        for item in results
    ]

    with csv_path.open("w", encoding="utf-8-sig", newline="") as stream:
        writer = csv.writer(stream, delimiter=";")
        writer.writerow(headers)
        writer.writerows(rows)

    workbook = xlwt.Workbook(encoding="utf-8")
    sheet = workbook.add_sheet("Отчёт")
    header_style = xlwt.easyxf(
        "font: bold on, colour white; pattern: pattern solid, fore_colour dark_blue; "
        "align: vert centre, horiz centre; borders: bottom thin, top thin, left thin, right thin"
    )
    cell_style = xlwt.easyxf(
        "align: wrap on, vert top; borders: bottom thin, top thin, left thin, right thin"
    )
    for col, header in enumerate(headers):
        sheet.write(0, col, header, header_style)
    for row_index, values in enumerate(rows, 1):
        for col, value in enumerate(values):
            sheet.write(row_index, col, value, cell_style)
    widths = [12, 30, 22, 40, 38, 18, 70, 60, 16, 70]
    for index, width in enumerate(widths):
        sheet.col(index).width = min(width * 256, 65535)
    sheet.panes_frozen = True
    sheet.horz_split_pos = 1
    workbook.save(str(xls_path))
    return csv_path, xls_path


def print_project_parse(project: ProjectData) -> None:
    print(f"\nФайл: {project.source_file.name}")
    print(f"Проект: {project.title}")
    print(f"Регион: {project.region or 'не найден'}")
    print(f"Масштаб: {project.scale or 'не найден'}")
    print(f"Сроки: {project.start_date or '?'} — {project.end_date or '?'}")
    print(f"Мероприятий распознано: {len(project.calendar)}")
    print(f"Участников по результатам: {project.participants_count or 'не найдено'}")
    print(f"Позиций софинансирования: {len(project.cofinancing)}")
    print(f"Позиций сметы: {len(project.expenses)}")
    for warning in project.parser_warnings:
        print(f"Предупреждение: {warning}")


async def run_download(
    browser: Browser,
    rows: list[SourceAccount],
    settings: Settings,
    interactive_project_selection: bool,
) -> list[TransferResult]:
    results: list[TransferResult] = []
    for position, account in enumerate(rows, 1):
        hint = f" — «{account.project_hint}»" if account.project_hint else ""
        print(f"\n[скачивание {position}/{len(rows)}] {account.fio}{hint}")
        result = await process_download_account(
            browser, account, settings, interactive_project_selection
        )
        results.append(result)
        print(f"Статус: {result.status}")
        if result.project_name:
            print(f"Проект: {result.project_name}")
        if result.word_file:
            print(f"Файл: {result.word_file}")
        if result.error:
            print(f"Ошибка: {result.error}")
        if settings.stop_on_error and result.status == "ошибка":
            break
    return results


async def run_upload(
    browser: Browser,
    rows: list[TargetAccount],
    settings: Settings,
) -> list[TransferResult]:
    results: list[TransferResult] = []
    for position, account in enumerate(rows, 1):
        print(
            f"\n[перенос {position}/{len(rows)}] {account.fio} — "
            f"{account.word_reference}"
        )
        result = await process_upload_account(browser, account, settings)
        results.append(result)
        print(f"Статус: {result.status}")
        print(f"Заполнено полей: {result.fields_filled}")
        if result.error:
            print(f"Ошибка: {result.error}")
        if settings.stop_on_error and result.status == "ошибка":
            break
    return results


async def async_main() -> int:
    args = parse_args()
    config_path = Path(args.config).resolve()
    settings = load_settings(config_path)
    mode = args.mode or choose_menu_mode()

    settings.downloads_dir.mkdir(parents=True, exist_ok=True)
    settings.logs_dir.mkdir(parents=True, exist_ok=True)
    settings.screenshots_dir.mkdir(parents=True, exist_ok=True)

    source_rows: list[SourceAccount] = []
    target_rows: list[TargetAccount] = []
    if mode in {MODE_DOWNLOAD, MODE_ALL}:
        source_rows = filter_rows(
            read_source_accounts(settings.source_accounts_file), args.row, args.limit
        )
        if (
            mode == MODE_DOWNLOAD
            and settings.interactive_account_selection
            and not args.all_accounts
            and not args.row
            and not args.limit
        ):
            source_rows = choose_source_accounts(source_rows)
    if mode in {MODE_UPLOAD, MODE_ALL, MODE_PARSE}:
        target_rows = filter_rows(
            read_target_accounts(settings.target_accounts_file), args.row, args.limit
        )

    if mode == MODE_PARSE:
        failed = 0
        for account in target_rows:
            try:
                word_file = resolve_word_file(account.word_reference, settings)
                project = parse_project_docx(word_file, account.title_override)
                print_project_parse(project)
            except Exception as exc:
                failed += 1
                print(f"\nСтрока {account.row_number}: ОШИБКА — {exc}")
        return 0 if failed == 0 else 1

    results: list[TransferResult] = []
    async with async_playwright() as playwright:
        browser = await playwright.chromium.launch(
            headless=settings.headless,
            slow_mo=settings.slow_mo_ms,
            args=[
                "--disable-background-networking",
                "--disable-component-update",
                "--disable-default-apps",
                "--disable-extensions",
                "--disable-sync",
                "--metrics-recording-only",
                "--no-first-run",
            ],
        )
        try:
            if source_rows:
                results.extend(
                    await run_download(
                        browser,
                        source_rows,
                        settings,
                        settings.interactive_project_selection
                        and not args.auto_select,
                    )
                )
            if target_rows:
                results.extend(await run_upload(browser, target_rows, settings))
        finally:
            await browser.close()

    csv_report, xls_report = write_reports(results, settings)
    print("\nОбработка завершена")
    print(f"Отчёт XLS: {xls_report}")
    print(f"Отчёт CSV: {csv_report}")
    return 0 if all(item.status != "ошибка" for item in results) else 1


def main() -> None:
    try:
        raise SystemExit(asyncio.run(async_main()))
    except AutomationError as exc:
        print(f"ОШИБКА НАСТРОЙКИ: {exc}", file=sys.stderr)
        raise SystemExit(2)
    except KeyboardInterrupt:
        print("\nОстановлено пользователем", file=sys.stderr)
        raise SystemExit(130)


if __name__ == "__main__":
    main()
